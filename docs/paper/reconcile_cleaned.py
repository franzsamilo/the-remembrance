"""Reconcile the hand-cleaned docx (figures/formatting) with the verified content.

Pass 1 (4 edits): integrated the honest §3.3.3 campaign-query description, the
live recovered grounding (0.984), the §5.8.1 nine-seed KPI append, and the live
+75%/+423% uplift in the abstract + conclusion.

Pass 2 (+7 edits, H1 uplift reconciliation): the abstract/conclusion now headline
the LIVE 18-query evaluation (§5.8.1 -> +75%/+423%), but Chapter 5's H1 confirmation
reports the CAMPAIGN-graph 5-query probe set (§5.9 -> +45%/+204%). Both numbers are
real and both are kept; Pass 2 labels every uplift figure by which evaluation
produced it and bridges §5.9/§5.10 to the final live evaluation, so the summary no
longer silently contradicts the abstract. No true number is deleted.

Idempotent by construction: always reads the pristine 6.11_cleaned INPUT and writes
a fresh 6.12_cleaned OUTPUT, so re-running reproduces the same result and never
double-applies. Figures/tables/formatting in the hand-cleaned base are preserved.

Input : Project_Study_Report__The_Remembrance_6_11_cleaned.docx
Output: Project_Study_Report__The_Remembrance_6_12_cleaned.docx
"""
from __future__ import annotations

import os
import sys
from docx import Document

try:  # keep §, —, ±, ≥ from crashing the Windows cp1252 console
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(ROOT, "Project_Study_Report__The_Remembrance_6_11_cleaned.docx")
OUTPUT_DOCX = os.path.join(ROOT, "Project_Study_Report__The_Remembrance_6_12_cleaned.docx")

# (old, new, label). Applied in order to every body + table-cell paragraph.
EDITS = [
    # ----- Pass 1 -----
    # (1) §3.3.3 — describe the REAL general campaign queries, honestly.
    ("Third, the queries are deliberately corpus-aligned rather than gotcha-aligned. The five "
     "queries cover: (a) constitutional rights across decisions, (b) majority-opinion authorship "
     "attribution, (c) precedent citation and modification, (d) procedural standards for motions "
     "for reconsideration, and (e) intellectual property disputes. These represent open-ended "
     "legal-research questions a practitioner would pose, not adversarial probes constructed to "
     "favor any particular architectural outcome.",
     "Third, the queries are open-ended rather than gotcha-aligned. The five campaign queries are "
     "general, domain-agnostic corpus probes — the principal findings, the main contributors, the "
     "methods applied, the principal results, and the datasets and concepts discussed — chosen to "
     "exercise retrieval and synthesis end-to-end without being constructed to favour any "
     "particular architectural outcome. The final live-system evaluation (§5.8.1) replaces these "
     "with eighteen corpus-aligned intellectual-property questions specific to this corpus (the "
     "dominancy and holistic tests, unfair competition, copyright scope, collective-rights "
     "enforcement, and IPOPHL administrative procedure).",
     "§3.3.3 query desc"),
    # (2) abstract uplift -> live
    ("baseline is +45% Grounding and +204% Faithfulness",
     "baseline is +75% Grounding and +423% Faithfulness",
     "abstract uplift -> live"),
    # (3) conclusion uplift -> live (leaves campaign-progression +45% spots intact)
    ("+45% Grounding and +204% Faithfulness uplift of the full-stack",
     "+75% Grounding and +423% Faithfulness uplift of the full-stack",
     "conclusion uplift -> live"),
    # (4) §5.8.1 — append 9-seed KPIs + the committed second grounding run
    ("and the structural KPIs (AUC-ROC, MRR) hold on the rebuilt graph.",
     "and the structural KPIs reproduce on the rebuilt graph (nine-seed mean: AUC-ROC 0.986, "
     "MRR 0.959, 12/12 PASS). Against the prompt-only chunk-RAG baseline on the same eighteen-query "
     "set, the integrity layer's uplift is +75% Grounding and +423% Faithfulness (baseline "
     "0.561 / 0.188); a second committed five-run evaluation reproduced grounding at 0.993 ± 0.011 "
     "(four of five runs ≥ 0.98), confirming the result is not a single-run artifact. The passing "
     "artifacts are committed to the repository (multi_seed_confirm_current.json, "
     "reroll_grounding_confirm.json, reroll_eval_commit.json).",
     "§5.8.1 KPI append"),

    # ----- Pass 2: H1 uplift reconciliation (campaign §5.9 vs live §5.8.1) -----
    # (5) §5.9 progression: label the campaign figures and bridge to the live eval.
    ("remains substantially elevated at Run 8 (+204%).",
     "remains substantially elevated at Run 8 (+204%). These campaign-graph figures are measured "
     "on the five-query probe set; the final live 18-query intellectual-property evaluation "
     "(§5.8.1) reports +75% Grounding and +423% Faithfulness over the same prompt-only baseline.",
     "§5.9 campaign label + bridge"),
    # (6) §5.10 summary sentence: cite BOTH evals, each labelled.
    ("H1 is qualitatively confirmed by the +45% Grounding and +204% Faithfulness uplift over "
     "standard chunk-RAG baseline.",
     "H1 is confirmed by a consistent integrity-layer uplift over the chunk-RAG baseline: "
     "+45% Grounding / +204% Faithfulness on the campaign-graph probe set (§5.9) and "
     "+75% Grounding / +423% Faithfulness on the final live 18-query intellectual-property "
     "evaluation (§5.8.1).",
     "§5.10 summary -> both labelled"),
    # (7) Table 5.12 H1 cell (comma form is unique to the cell): show both, labelled.
    ("+45% Grounding, +204% Faithfulness",
     "+45%/+204% (campaign §5.9); +75%/+423% (live §5.8.1)",
     "Table 5.12 H1 cell -> both"),
    # (8) Ch3 operationalization: tie the §5.9 campaign result to the live result.
    ("Result (§5.9): full-stack achieves +45% Grounding and +204% Faithfulness uplift, confirming H1.",
     "Result (§5.9): full-stack achieves +45% Grounding and +204% Faithfulness uplift on the "
     "campaign-graph probe set, confirming H1; the final live 18-query evaluation (§5.8.1) yields "
     "+75% Grounding and +423% Faithfulness.",
     "§3.3.x H1 result + live"),
    # (9) Abstract: tag the headline live figure (runs after edit 2 created +75%).
    ("+75% Grounding and +423% Faithfulness, confirming the Topological Correlation hypothesis (H1).",
     "+75% Grounding and +423% Faithfulness on the final live 18-query evaluation, confirming the "
     "Topological Correlation hypothesis (H1).",
     "abstract live tag"),
    # (10) Conclusion: tag the headline live figure (runs after edit 3 created +75%).
    ("+75% Grounding and +423% Faithfulness uplift of the full-stack",
     "+75% Grounding and +423% Faithfulness (final live 18-query evaluation) uplift of the full-stack",
     "conclusion live tag"),
]

# Exact-match caption labels (== avoids touching List-of-Tables / TOC entries,
# which carry trailing tabs + page numbers and so never match exactly).
CAP_OLD = "Table 5.11: Full-Stack vs Prompt-Only Comparison"
CAP_SUFFIX = " (Campaign-Graph Probe Set)"


def replace_once(par, old, new):
    runs = par.runs
    full = "".join(r.text for r in runs)
    idx = full.find(old)
    if idx == -1:
        return 0
    end = idx + len(old)
    pos = 0; ranges = []
    for r in runs:
        ranges.append((pos, pos + len(r.text), r)); pos += len(r.text)
    inserted = False
    for rs, re_, r in ranges:
        if re_ <= idx or rs >= end:
            continue
        ls, le = max(idx, rs) - rs, min(end, re_) - rs
        seg = r.text
        r.text = (seg[:ls] + new + seg[le:]) if not inserted else (seg[:ls] + seg[le:])
        inserted = True
    return 1


def main():
    if not os.path.exists(INPUT_DOCX):
        print(f"FATAL no input: {INPUT_DOCX}"); return 1
    print(f"OPEN {os.path.basename(INPUT_DOCX)}", flush=True)
    doc = Document(INPUT_DOCX)
    paras = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                paras.extend(c.paragraphs)

    applied = 0
    for i, (old, new, label) in enumerate(EDITS, 1):
        hit = sum(replace_once(p, old, new) for p in paras)
        print(f"  EDIT {i:2} [{label}]: {'APPLIED x%d' % hit if hit else 'SKIPPED (target not found)'}", flush=True)
        applied += hit

    # caption label (exact-match, append to last run to preserve formatting)
    cap_hits = 0
    for p in paras:
        if p.text.strip() == CAP_OLD and p.runs and not p.text.strip().endswith("Probe Set)"):
            p.runs[-1].text = p.runs[-1].text + CAP_SUFFIX
            cap_hits += 1
    print(f"  CAP    [Table 5.11 caption label]: {'APPLIED x%d' % cap_hits if cap_hits else 'SKIPPED'}", flush=True)

    # verify
    joined = "\n".join(p.text for p in doc.paragraphs)
    cell_join = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    all_txt = joined + "\n" + cell_join
    want = {
        "general, domain-agnostic corpus probes": "§3.3.3 honest rewrite",
        "on the final live 18-query evaluation, confirming the": "abstract tagged live",
        "(final live 18-query evaluation) uplift of the full-stack": "conclusion tagged live",
        "nine-seed mean: AUC-ROC 0.986": "§5.8.1 KPIs appended",
        "These campaign-graph figures are measured on the five-query probe set": "§5.9 campaign label + bridge",
        "+45% Grounding / +204% Faithfulness on the campaign-graph probe set (§5.9) and": "§5.10 both labelled",
        "+45%/+204% (campaign §5.9); +75%/+423% (live §5.8.1)": "Table 5.12 H1 both",
        "campaign-graph probe set, confirming H1; the final live 18-query": "§3.3.x H1 + live",
        "(Campaign-Graph Probe Set)": "Table 5.11 caption labelled",
        # legitimate campaign numbers that MUST survive untouched:
        "+10% at Run 1, +32% at Run 6, +45% at Run 8": "campaign progression intact",
        "+0.306 (+45%)": "Table 5.11 campaign delta intact",
    }
    absent = {
        "constitutional rights across these decisions": "OLD §3.3.3 gone",
        "qualitatively confirmed by the +45% Grounding and +204% Faithfulness uplift over standard chunk-RAG baseline": "OLD unlabelled §5.10 gone",
    }
    print("VERIFY present:", flush=True)
    ok = True
    for tok, desc in want.items():
        p = tok in all_txt
        ok &= p
        print(f"  {'YES' if p else 'NO ':3} {desc}", flush=True)
    print("VERIFY absent:", flush=True)
    for tok, desc in absent.items():
        a = tok not in all_txt
        ok &= a
        print(f"  {'YES' if a else 'NO ':3} {desc}", flush=True)

    imgs = sum(1 for r in doc.part.rels.values() if "image" in r.reltype)
    doc.save(OUTPUT_DOCX)
    print(f"SAVED {os.path.basename(OUTPUT_DOCX)}  (edits applied: {applied}/{len(EDITS)}, caption: {cap_hits})", flush=True)
    print(f"tables={len(doc.tables)} paras={len(doc.paragraphs)} images={imgs} (figures preserved)", flush=True)
    print(f"OVERALL: {'PASS' if ok else 'FAIL — review above'}", flush=True)
    return 0 if ok else 2


if __name__ == "__main__":
    raise SystemExit(main())
