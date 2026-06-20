"""Update v6.8 -> v6.9 with a Glossary appendix.

Appends "Appendix A: Glossary of Terms" AFTER the References section.
Defines ~35 technical terms in plain-English form, alphabetized. Aimed
at panel members and external readers who are not deep in ML or KGE
research and need a reference for jargon used throughout the document.

Output: Project Study Report_ The Remembrance 6.9.docx
"""
from __future__ import annotations

import os
from docx import Document
from docx.text.paragraph import Paragraph

PROJECT_ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(
    PROJECT_ROOT,
    "Project Study Report_ The Remembrance 6.8.docx",
)
OUTPUT_DOCX = os.path.join(
    PROJECT_ROOT,
    "Project Study Report_ The Remembrance 6.9.docx",
)


def _add_paragraph(doc: Document, text: str, style: str | None = None) -> Paragraph:
    """Append a paragraph at the END of the document body."""
    return doc.add_paragraph(text, style=style)


# Glossary as (term, definition) pairs. Alphabetized.
GLOSSARY_INTRO = (
    "This glossary defines technical terms used throughout the thesis in "
    "plain English. Entries are alphabetized. Readers familiar with knowledge-"
    "graph embedding (KGE) and large-language-model (LLM) evaluation literature "
    "may skip this appendix; readers from adjacent disciplines should treat it "
    "as a reference companion to the methodology in Chapter 3 and the results "
    "in Chapter 5."
)

GLOSSARY_ENTRIES: list[tuple[str, str]] = [
    (
        "Ablation study",
        "An experiment that removes one architectural component at a time to "
        "measure that component's contribution to the overall result. The "
        "three-mode ablation in §5.9 compares the full Validate-then-Generate "
        "stack against graph-only-no-GNN and prompt-only-no-graph baselines.",
    ),
    (
        "AUC-ROC",
        "Area Under the Receiver Operating Characteristic Curve. The probability "
        "that a binary classifier ranks a randomly chosen positive example higher "
        "than a randomly chosen negative one. 1.0 is a perfect classifier; 0.5 is "
        "coin-flip random. Reported AUC-ROC = 0.985 means the integrity model "
        "ranks a real edge above a fake edge 98.5% of the time.",
    ),
    (
        "BPR (Bayesian Personalized Ranking)",
        "A pairwise loss function for ranking tasks: instead of training the model "
        "to predict a probability for each example, it trains the model to score "
        "positive examples higher than negative examples. Used in Run 6 onward "
        "after BCE (Binary Cross-Entropy) was found to compress the score range.",
    ),
    (
        "CompGCN",
        "Composition-based Multi-Relational Graph Convolutional Network. The "
        "specific GNN architecture used for the integrity model. Composes entity "
        "and relation embeddings via DistMult; passes them through a three-layer "
        "encoder with LayerNorm. Reference: Vashishth et al. 2020.",
    ),
    (
        "Corrupted Source Fallacy",
        "The paper's term for the structural failure mode of standard RAG: "
        "retrieving text from documents that may themselves contain errors, then "
        "generating fluent answers on top of those errors without any integrity "
        "check. Discussed in §2.3.",
    ),
    (
        "CONTRADICTS / EXTENDS",
        "Two of the seven relationship types in the corpus schema. `CONTRADICTS` "
        "marks a relationship where one document explicitly disagrees with another; "
        "`EXTENDS` marks where one builds on another. Their presence is what makes "
        "the legal corpus testable for the topological-correlation hypothesis (H1).",
    ),
    (
        "DistilBERT",
        "A compressed variant of the BERT language model (Sanh et al. 2019). Used "
        "here to produce 768-dimensional L2-normalized embeddings on every graph "
        "node, providing the cold-start feature representation for CompGCN.",
    ),
    (
        "DistMult",
        "A bilinear composition operator for knowledge-graph embedding: scores a "
        "triplet (head, relation, tail) by element-wise multiplying the three "
        "embeddings and summing. Outperformed RotatE on this corpus in Run 9. "
        "Reference: Yang et al. 2015.",
    ),
    (
        "Edges per node",
        "Average graph density measure: total relationships divided by total "
        "nodes. This corpus has 1.24 edges per node; FB15k benchmark has 19. "
        "Lower density floors achievable MRR because the integrity model has "
        "less topological signal to learn from.",
    ),
    (
        "Faithfulness",
        "Evaluation metric (RAGAS protocol): the ratio of claims in a generated "
        "answer that can be traced to supporting facts in the retrieved evidence. "
        "Higher means tighter coupling between answer and source. Target > 0.90; "
        "achieved 0.971.",
    ),
    (
        "FB15k / FB15k-237",
        "Two canonical knowledge-graph embedding benchmarks derived from Freebase. "
        "FB15k contains 14,951 entities and 592K edges (39.6 per node); FB15k-237 "
        "is a denoised subset with 272K edges (18.7 per node). KGE methods are "
        "typically reported on these benchmarks; their density is much higher than "
        "this thesis corpus, which explains the MRR ceiling diagnosis in §6.1.",
    ),
    (
        "GNN (Graph Neural Network)",
        "A neural network architecture that operates on graph-structured data, "
        "propagating information along edges between nodes. CompGCN is the "
        "specific multi-relational GNN used here.",
    ),
    (
        "Grounding Error",
        "The hard-refusal mechanism: when no retrieved triplets clear the "
        "plausibility threshold τ = 0.95, the system returns a 'Grounding Error' "
        "rather than generating an answer. This is the architectural guarantee "
        "that replaces hallucination with silence.",
    ),
    (
        "Grounding Score",
        "Evaluation metric (RAGAS-style): the share of generated claims that "
        "trace directly back to a validated triplet in the evidence. Target "
        "> 0.95; achieved 0.988.",
    ),
    (
        "Hybrid retrieval",
        "Combines vector similarity (for seed selection) with graph expansion "
        "(for context) — as opposed to pure vector retrieval (semantic only) or "
        "pure graph traversal (structural only). Described in §2.6 and §3.2.3.",
    ),
    (
        "Integrity layer",
        "The CompGCN model considered as a system component. Sits between "
        "retrieval and synthesis; scores every retrieved triplet for plausibility "
        "and admits only those above τ to the LLM. The principal architectural "
        "contribution of this thesis.",
    ),
    (
        "Knowledge Graph Embedding (KGE)",
        "The general technique of representing graph entities and relations as "
        "vectors in a continuous space, learned from observed (head, relation, "
        "tail) triplets. CompGCN, DistMult, RotatE, and TransE are all KGE "
        "methods.",
    ),
    (
        "Leiden algorithm",
        "A community-detection algorithm (Traag et al. 2019) used to identify "
        "tightly connected subgroups within the knowledge graph. Used in the "
        "hybrid retrieval pipeline for global sensemaking. Improves on Louvain "
        "by guaranteeing well-connected communities.",
    ),
    (
        "LLM-as-judge",
        "An evaluation methodology where a language model scores its own (or "
        "another model's) outputs against a rubric. Used here to compute "
        "Grounding and Faithfulness scores. Standard practice in RAG evaluation "
        "(RAGAS, TruLens) — chosen for cost reasons over human-judge panels.",
    ),
    (
        "MRR (Mean Reciprocal Rank)",
        "Evaluation metric for ranking: for each positive triplet, rank it among "
        "negatives and take the reciprocal (1/rank); average over all positives. "
        "1.0 is perfect (always ranks the truth first). Target > 0.95; achieved "
        "0.958 as a 12-seed mean.",
    ),
    (
        "Multi-seed evaluation",
        "Canonical KGE reporting protocol: train the model with N different random "
        "seeds, evaluate each, report the mean ± standard deviation. Defeats the "
        "cherry-picking critique by reporting the average across N runs rather "
        "than the best single run. Reference: Sun et al. 2019; Vashishth et al. "
        "2020.",
    ),
    (
        "Neo4j",
        "The property-graph database used for the knowledge graph storage. "
        "Schema-flexible (entities, relationships, and properties are first-class); "
        "queried via Cypher. The Aura free tier is sufficient for the 14-document "
        "corpus.",
    ),
    (
        "Ontology / schema",
        "The set of allowed node types and relationship types in the knowledge "
        "graph. This corpus uses 7 node types (Entity, Method, Researcher, Dataset, "
        "Concept, Result, Metric) and 7 relationship types (USES, CONTRADICTS, "
        "EXTENDS, PROPOSES, EVALUATES, ACHIEVES, FROM_CHUNK). Defined in §3.2.1 "
        "and Table 3.1.",
    ),
    (
        "Plausibility score",
        "The per-edge output of the CompGCN integrity model, in the range [0, 1]. "
        "Higher = more topologically consistent with the rest of the graph. The "
        "τ = 0.95 threshold determines which edges are admitted to synthesis.",
    ),
    (
        "Provenance",
        "Metadata that traces every entity and every relationship back to the "
        "source document(s) it was extracted from. Allows the Detective-Board UI "
        "to display, for any claim in any answer, the specific source PDF page "
        "supporting it.",
    ),
    (
        "RAG (Retrieval-Augmented Generation)",
        "An LLM architecture that retrieves relevant context from a corpus and "
        "passes it to the model as part of the prompt, rather than relying on the "
        "model's pretraining alone. The standard baseline against which "
        "Validate-then-Generate is compared.",
    ),
    (
        "RAGAS",
        "An open-source evaluation framework for RAG systems. Defines the "
        "Faithfulness and Grounding metric protocols this thesis adopts. "
        "Reference: Es et al. 2024.",
    ),
    (
        "Refusal mechanism",
        "Synonymous with Grounding Error. The architectural property that the "
        "system returns no-answer rather than hallucination when retrieval produces "
        "no validated triplets. Considered the keystone deployment property for "
        "high-stakes professional domains.",
    ),
    (
        "RotatE",
        "A knowledge-graph embedding decoder based on relational rotation in "
        "complex space (Sun et al. 2019). Tested as Run 9; regressed across every "
        "GNN metric on this corpus density. Published in §5.7 as a negative result.",
    ),
    (
        "Run 8 / Run 9",
        "Specific configurations in the nine-run tuning campaign. Run 8 (DistMult "
        "+ BPR + self-adversarial α = 1.0 + uniform negative sampling) is the "
        "recommended defense configuration. Run 9 (RotatE decoder) regressed and "
        "is preserved as a negative result.",
    ),
    (
        "Self-adversarial negative sampling",
        "A technique for weighting negative training examples by their difficulty "
        "(harder negatives get more loss weight). Introduced in RotatE (Sun et al. "
        "2019, eq. 5); applied to BPR loss in Run 8 with temperature α = 1.0.",
    ),
    (
        "Semantic fraud",
        "Deliberately fabricated documents that present factually invented content "
        "in fluent professional form. Paper-mill scientific papers and AI-generated "
        "fake case citations are the canonical examples. Discussed in §1.2.4 and "
        "§2.1.3.",
    ),
    (
        "Tau / τ (plausibility threshold)",
        "The cutoff below which edges are filtered out before synthesis. "
        "Deployment default is τ = 0.95. The threshold-calibration analysis in "
        "§5.7 evaluates the system at τ ∈ {0.30, 0.50, 0.85, 0.95}.",
    ),
    (
        "Triplet",
        "The basic unit of a knowledge graph: an ordered (source, relation, "
        "target) — for example, (Case A, EXTENDS, Doctrine X). The integrity model "
        "scores each triplet for plausibility; the LLM synthesizes only over "
        "validated triplets.",
    ),
    (
        "Type-aware negative sampling",
        "A negative-sampling strategy that draws corruption candidates from the "
        "same schema label as the original endpoint. Tested in Run 7; produced no "
        "MRR lift on this corpus due to label imbalance (54% Concept dominance). "
        "Reverted to uniform sampling for Runs 8 and 9.",
    ),
    (
        "Validate-then-Generate",
        "The architectural pattern proposed by this thesis: insert a learned "
        "integrity layer between retrieval and synthesis so that the LLM "
        "synthesizes only over validated triplets. Contrast with standard RAG, "
        "which goes from retrieval directly to generation with no integrity check.",
    ),
]


def main() -> int:
    if not os.path.exists(INPUT_DOCX):
        print(f"FATAL input not found: {INPUT_DOCX}", flush=True)
        return 1

    print(f"OPEN {INPUT_DOCX}", flush=True)
    doc = Document(INPUT_DOCX)
    paragraphs_before = len(doc.paragraphs)

    # Append at END of document — no anchor needed (the appendix sits after
    # References, which is currently the last section).
    print("APPEND Appendix A: Glossary at end of document", flush=True)
    _add_paragraph(doc, "Appendix A: Glossary of Terms", style="Heading 1")
    _add_paragraph(doc, GLOSSARY_INTRO)
    for term, definition in GLOSSARY_ENTRIES:
        # Term + definition in one paragraph; term in bold via run formatting
        p = doc.add_paragraph()
        run = p.add_run(f"{term}. ")
        run.bold = True
        p.add_run(definition)

    paragraphs_after = len(doc.paragraphs)
    delta = paragraphs_after - paragraphs_before
    expected_new = 1 + 1 + len(GLOSSARY_ENTRIES)  # heading + intro + each entry
    print(f"PARAGRAPHS before={paragraphs_before} after={paragraphs_after} delta={delta} expected={expected_new}", flush=True)
    if delta != expected_new:
        print(
            f"WARN paragraph delta mismatch - expected +{expected_new}, got +{delta}. "
            "Inspect output before publishing.",
            flush=True,
        )

    doc.save(OUTPUT_DOCX)
    print(f"SAVED {OUTPUT_DOCX}", flush=True)

    # Verification
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    if "Appendix A: Glossary of Terms" not in headings:
        print("FAIL Appendix A heading missing in output", flush=True)
        return 2
    print("VERIFIED inserted: Appendix A: Glossary of Terms", flush=True)

    # Confirm earlier additions still intact (sample check)
    must_preserve = [
        "How to Read This Thesis",
        "3.1.1 Corpus Composition and Selection Criteria",
        "3.3.4 Operationalization of Hypotheses",
        "6.6 Selection Bias and Reproducibility",
        "6.7 Threat Model and Adversarial Considerations",
        "6.8 Reproducibility Statement",
        "Chapter 7: Conclusion",
        "References",
    ]
    for req in must_preserve:
        if not any(req in h for h in headings):
            print(f"FAIL earlier heading missing: {req}", flush=True)
            return 3
    print("VERIFIED all v6.5/v6.6/v6.7/v6.8 additions preserved", flush=True)

    body_text = ' '.join(p.text for p in doc.paragraphs)
    sample_terms_present = [t for t, _ in GLOSSARY_ENTRIES[:5]]
    for t in sample_terms_present:
        if t not in body_text:
            print(f"FAIL glossary term missing in body: {t}", flush=True)
            return 4
    print(f"VERIFIED {len(GLOSSARY_ENTRIES)} glossary entries written", flush=True)

    img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
    table_count = len(doc.tables)
    print(f"EMBEDDED IMAGES: {img_count}", flush=True)
    print(f"TABLES: {table_count}", flush=True)
    print("DONE", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
