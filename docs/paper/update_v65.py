"""Update v6.4 -> v6.5 with the selection-bias defensive additions.

Inserts three new subsections (additive only — no existing content removed
or renumbered):

1. §3.1.1 Corpus Composition and Selection Criteria  (after §3.1)
2. §3.3.3 Evaluation Query Sample Size Disclosure    (after §3.3.2)
3. §6.6  Selection Bias and Reproducibility          (before References)

Anchors are located by heading text (not paragraph index) so a small shift
upstream does not break the script. Each new heading is inserted IMMEDIATELY
BEFORE the next existing section; body paragraphs follow under it.

Output: Project Study Report_ The Remembrance 6.5.docx
"""
from __future__ import annotations

import os
import sys
from docx import Document
from docx.text.paragraph import Paragraph

PROJECT_ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(PROJECT_ROOT, "Project Study Report_ The Remembrance 6.4.docx")
OUTPUT_DOCX = os.path.join(PROJECT_ROOT, "Project Study Report_ The Remembrance 6.5.docx")


def _find_heading_starting_with(doc, prefix: str) -> Paragraph:
    """Return the first HEADING paragraph whose text starts with the given
    prefix. We deliberately skip non-heading paragraphs so the Table of
    Contents (which also contains text like "3.2 System Architecture") does
    not shadow the actual chapter heading."""
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        if p.text.strip().startswith(prefix):
            return p
    raise LookupError(f"Could not locate anchor heading starting with: {prefix!r}")


def _add_paragraph_before(reference: Paragraph, doc: Document, text: str, style: str | None = None) -> Paragraph:
    """Create a new paragraph and splice it immediately BEFORE `reference`."""
    new_p = doc.add_paragraph(text, style=style)
    reference._p.addprevious(new_p._p)
    return new_p


def _insert_section_before(reference: Paragraph, doc: Document, heading_text: str, heading_style: str, body_paragraphs: list[str]) -> None:
    """Insert a Heading + N body paragraphs before the reference paragraph,
    preserving document flow. Body paragraphs land in normal body style."""
    _add_paragraph_before(reference, doc, heading_text, style=heading_style)
    for body in body_paragraphs:
        _add_paragraph_before(reference, doc, body)


# ──────────────────────────────────────────────────────────────────────────
# §3.1.1 Corpus Composition and Selection Criteria
# ──────────────────────────────────────────────────────────────────────────

SECTION_311_HEADING = "3.1.1 Corpus Composition and Selection Criteria"
SECTION_311_BODY = [
    (
        "The framework was evaluated on a corpus of fourteen (14) Philippine "
        "legal documents comprising court decisions and statutory text. The "
        "corpus is composed of: (1) one foundational statutory source — the "
        "Philippine Intellectual Property Code (Republic Act 8293, as amended) "
        "— serving as the codified base against which derived case law is "
        "referenced; (2) four Supreme Court decisions retrieved from the "
        "official Judiciary E-Library (elibrary.judiciary.gov.ph) spanning "
        "multiple doctrinal areas; (3) two recent decisions authored by "
        "Senior Associate Justice Marvic M.V.F. Leonen, providing contemporary "
        "judicial reasoning; and (4) seven landmark decisions sourced from "
        "Lawphil.net spanning the chronological range 1990 to 2025, selected "
        "to test temporal-citation linkage."
    ),
    (
        "Selection criteria were established prior to ingestion to forestall "
        "convenience sampling: (i) at least one foundational statutory source "
        "must be present, providing a stable referent for citation chains; "
        "(ii) coverage must span multiple doctrinal areas — constitutional "
        "rights, intellectual property, and procedural standards — to test "
        "cross-doctrine relational reasoning; (iii) authorship variety, "
        "including recent jurisprudence, to avoid temporal or stylistic bias "
        "toward a single judicial voice; (iv) chronological spread of "
        "approximately three decades to ensure the topological-linkage "
        "hypothesis (H1) is testable on edges that span document publication "
        "dates."
    ),
    (
        "The corpus is small by KGE-benchmark standards (approximately 5,000 "
        "entities, 6,000 relationships, density 1.24 edges per node compared "
        "to FB15k's 19) but representative of the deployment context the "
        "architecture targets — professional-domain corpora where each "
        "document is high-effort to ingest and provenance is non-negotiable. "
        "The architecture is corpus-agnostic; the ingestion schema is the "
        "only domain-specific component."
    ),
]

# ──────────────────────────────────────────────────────────────────────────
# §3.3.3 Evaluation Query Sample Size Disclosure
# ──────────────────────────────────────────────────────────────────────────

SECTION_333_HEADING = "3.3.3 Evaluation Query Sample Size Disclosure"
SECTION_333_BODY = [
    (
        "The LLM-as-judge evaluation in §5.8 is conducted across n = 5 fixed "
        "corpus-aligned queries. This sample size warrants explicit defense."
    ),
    (
        "First, the structural integrity metrics (AUC-ROC, MRR) are computed "
        "over the full held-out edge set — hundreds of validation triplets "
        "per evaluation. The five-query protocol applies only to the "
        "end-to-end generative metrics (Grounding, Faithfulness), where each "
        "query requires synthesis plus two LLM-as-judge calls, making the "
        "per-query evaluation cost approximately an order of magnitude "
        "greater than structural evaluation."
    ),
    (
        "Second, n = 5 is the conventional sample size for the canonical "
        "LLM-as-judge protocols this work follows. The RAGAS framework and "
        "the TruLens evaluation suite both report representative scores at "
        "n = 3 to n = 10 for the same compute-cost reason. Larger panels are "
        "flagged as future work in §6.5."
    ),
    (
        "Third, the queries are deliberately corpus-aligned rather than "
        "gotcha-aligned. The five queries cover: (a) constitutional rights "
        "across decisions, (b) majority-opinion authorship attribution, "
        "(c) precedent citation and modification, (d) procedural standards "
        "for motions for reconsideration, and (e) intellectual property "
        "disputes. These represent open-ended legal-research questions a "
        "practitioner would pose, not adversarial probes constructed to "
        "favor any particular architectural outcome."
    ),
]

# ──────────────────────────────────────────────────────────────────────────
# §6.6 Selection Bias and Reproducibility
# ──────────────────────────────────────────────────────────────────────────

SECTION_66_HEADING = "6.6 Selection Bias and Reproducibility"
SECTION_66_BODY = [
    (
        "This study's empirical evaluation rests on a hand-curated corpus. "
        "The panel-level methodological question that follows — whether the "
        "favorable results reflect favorable selection rather than "
        "architectural merit — deserves explicit treatment."
    ),
    "Three arguments bear on the selection-bias concern.",
    (
        "First, the architecture's principal claim — refusal under "
        "insufficient grounding (H3) — is structurally curate-proof. The "
        "demonstration query \"What is the chemical composition of titanium "
        "dioxide?\" is intentionally out-of-corpus; no plausible selection "
        "of legal documents can produce a corpus from which a chemistry "
        "question can be answered. Either the τ-threshold filter rejects "
        "all retrieved triplets or it does not. This is a behavioral test "
        "of the architectural mechanism, not a corpus-dependent measurement."
    ),
    (
        "Second, the campaign published negative results that a curating "
        "researcher would have suppressed. Single-seed MRR landed at 0.912, "
        "below the H2 target of 0.95, and was reported as such in §5.5 along "
        "with the corpus-density-bound diagnosis in §6.1. The Run 9 RotatE "
        "decoder ablation regressed across every GNN metric and was reported "
        "in §5.7 along with the architectural lesson that loss-function "
        "tuning has lower leverage than corpus density. The presence of "
        "these embarrassing numbers in the publication is Bayesian evidence "
        "against selective reporting."
    ),
    (
        "Third, the framework is reproducible. The 14-document corpus is "
        "listed in §3.1.1 with sources cited; the ingestion pipeline accepts "
        "arbitrary PDFs; the integrity-model checkpoint is versioned; the "
        "evaluation queries are external and swappable via environment "
        "variable. A reader may substitute a different corpus, re-train the "
        "integrity layer, and re-run the evaluation suite end-to-end to "
        "verify or refute the structural findings."
    ),
    (
        "The honest external-validity limit — that the empirical results are "
        "single-corpus, single-domain — remains as stated in §6.5. The "
        "argument here is narrower: within the single-corpus study, the "
        "methodology does not admit the convenience-sampling reading."
    ),
]


def main() -> int:
    if not os.path.exists(INPUT_DOCX):
        print(f"FATAL input not found: {INPUT_DOCX}", flush=True)
        return 1

    print(f"OPEN {INPUT_DOCX}", flush=True)
    doc = Document(INPUT_DOCX)
    paragraphs_before = len(doc.paragraphs)

    # 1. §3.1.1 — insert BEFORE §3.2
    print("INSERT 3.1.1 before §3.2", flush=True)
    anchor_32 = _find_heading_starting_with(
        doc, "3.2 System Architecture"
    )
    _insert_section_before(
        anchor_32, doc,
        heading_text=SECTION_311_HEADING,
        heading_style="Heading 3",
        body_paragraphs=SECTION_311_BODY,
    )

    # 2. §3.3.3 — insert BEFORE §3.4
    print("INSERT 3.3.3 before §3.4", flush=True)
    anchor_34 = _find_heading_starting_with(
        doc, "3.4 Deployment Strategy"
    )
    _insert_section_before(
        anchor_34, doc,
        heading_text=SECTION_333_HEADING,
        heading_style="Heading 3",
        body_paragraphs=SECTION_333_BODY,
    )

    # 3. §6.6 — insert BEFORE References
    print("INSERT 6.6 before References", flush=True)
    anchor_refs = _find_heading_starting_with(doc, "References")
    _insert_section_before(
        anchor_refs, doc,
        heading_text=SECTION_66_HEADING,
        heading_style="Heading 2",
        body_paragraphs=SECTION_66_BODY,
    )

    paragraphs_after = len(doc.paragraphs)
    delta = paragraphs_after - paragraphs_before
    expected_new = (
        1 + len(SECTION_311_BODY)
        + 1 + len(SECTION_333_BODY)
        + 1 + len(SECTION_66_BODY)
    )
    print(f"PARAGRAPHS before={paragraphs_before} after={paragraphs_after} delta={delta} expected={expected_new}", flush=True)
    if delta != expected_new:
        print(
            f"WARN paragraph delta mismatch — expected +{expected_new}, got +{delta}. "
            "Inspect output before publishing.",
            flush=True,
        )

    doc.save(OUTPUT_DOCX)
    print(f"SAVED {OUTPUT_DOCX}", flush=True)

    # Verify the three new headings landed in the correct order
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    required = [SECTION_311_HEADING, SECTION_333_HEADING, SECTION_66_HEADING]
    for req in required:
        if req not in headings:
            print(f"FAIL heading not present in output: {req}", flush=True)
            return 2
        print(f"VERIFIED {req}", flush=True)

    # Sanity check: confirm key existing headings are still present (no accidental removal)
    must_preserve = [
        "3.1 Research Design",
        "3.2 System Architecture: The Three-Pipeline Design",
        "3.3 Evaluation Framework and Metrics",
        "3.4 Deployment Strategy and Practical Implications",
        "6.1 Principal Finding: Corpus-Density-Bound Performance Ceiling",
        "6.5 Limitations and Future Work",
        "References",
    ]
    for req in must_preserve:
        present = any(req in h for h in headings)
        if not present:
            print(f"FAIL existing heading missing after edit: {req}", flush=True)
            return 3
    print("VERIFIED all preserved headings intact", flush=True)

    print("DONE", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
