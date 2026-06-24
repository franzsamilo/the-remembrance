"""Update v6.11 -> v6.12: Option B — make Table 5.10 and §3.3.3 describe the REAL
corpus-aligned legal queries that were actually run on the live system, replacing the
generic research-paper probes ("key findings", "researchers") that mismatched §3.3.3's
legal description. Per-query scores are pulled from the committed run_logs/reroll_eval_commit.json
(5-run mean, seed=6000 / exp=10, tau=0.95). Also fixes the Table 5.10 Mean-row bug
(0.984/0.983 -> the true mean of its own rows) and adds a §5.8.1 line noting a second
committed grounding run (0.993) reproduced the result.

Headline grounding/faith stay 0.984/0.983 (backed by the committed reroll_grounding_confirm
log); reroll_eval_commit.json is the second, confirming committed run.

Output: Project Study Report_ The Remembrance 6.12.docx
"""
from __future__ import annotations

import os
from docx import Document

ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(ROOT, "Project Study Report_ The Remembrance 6.11.docx")
OUTPUT_DOCX = os.path.join(ROOT, "Project Study Report_ The Remembrance 6.12.docx")

# Table 5.10 = docx table index 13. Rebuild rows r1..r6 (5 IP queries + Mean).
TABLE_510_IDX = 13
TABLE_510_ROWS = [
    (1, "How do the courts contrast the dominancy test with the holistic test?", "1.000", "1.000"),
    (2, "What is copyright infringement?", "1.000", "1.000"),
    (3, "What remedies are available for intellectual property infringement?", "1.000", "1.000"),
    (4, "How do the courts assess likelihood of confusion between competing trademarks?", "1.000", "1.000"),
    (5, "What is the role of the Intellectual Property Office of the Philippines?", "0.920", "1.000"),
    (6, "Mean", "0.984", "1.000"),
]

PROSE_EDITS = [
    # §3.3.3 query description -> the real legal queries
    ("The five queries cover: (a) constitutional rights across decisions, (b) majority-opinion "
     "authorship attribution, (c) precedent citation and modification, (d) procedural standards "
     "for motions for reconsideration, and (e) intellectual property disputes.",
     "The five queries cover core intellectual-property doctrines and institutions in the corpus: "
     "(a) the contrast between the dominancy and holistic tests for trademark infringement, "
     "(b) copyright infringement, (c) remedies for intellectual-property infringement, "
     "(d) likelihood-of-confusion assessment between competing marks, and (e) the role of the "
     "Intellectual Property Office of the Philippines (IPOPHL)."),
    # Table 5.10 caption — provenance (live system, IP queries, subset of 18)
    ("for the Run 8 recommended configuration, evaluated by Gemini as LLM-as-Judge across the "
     "five fixed evaluation queries from Section 3.3.2.",
     "for the recommended configuration on the recovered live system (§5.8.1), evaluated by Gemini "
     "as LLM-as-Judge across five corpus-aligned intellectual-property queries (a subset of the "
     "eighteen-query re-evaluation set)."),
    # Table 5.10 caption — which query fell short
    ("the fifth (datasets and concepts) fell below 1.00 on both metrics due to a partial coverage "
     "gap in the validated triplet set for that query.",
     "the fifth (the role of the IPOPHL) fell below 1.00 on Grounding due to a partial coverage "
     "gap in the validated triplet set for that broad institutional query."),
    # §5.8.1 — note the second committed grounding run
    ("expose ungrounded chunk-RAG generation more sharply.",
     "expose ungrounded chunk-RAG generation more sharply. A second independent five-run "
     "evaluation, persisted to the repository alongside the first, reproduced grounding at "
     "0.993 ± 0.011 (four of five runs ≥ 0.98), confirming the result is not a single-run artifact."),
]


def replace_in_paragraph(par, old, new):
    """Single-shot run-aware replace (all edits here are single-occurrence; a while
    loop would spin forever when `new` contains `old`, e.g. an appending edit)."""
    runs = par.runs
    full = "".join(r.text for r in runs)
    idx = full.find(old)
    if idx == -1:
        return 0
    end = idx + len(old)
    pos = 0; ranges = []
    for r in runs:
        ranges.append((pos, pos + len(r.text), r)); pos += len(r.text)
    inserted = False
    for rs, re_, r in ranges:
        if re_ <= idx or rs >= end:
            continue
        ls, le = max(idx, rs) - rs, min(end, re_) - rs
        seg = r.text
        r.text = (seg[:ls] + new + seg[le:]) if not inserted else (seg[:ls] + seg[le:])
        inserted = True
    return 1


def all_paras(doc):
    yield from doc.paragraphs
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                yield from c.paragraphs


def count_occ(doc, s):
    return sum(p.text.count(s) for p in all_paras(doc))


def set_cell(cell, text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def main():
    if not os.path.exists(INPUT_DOCX):
        print(f"FATAL no input: {INPUT_DOCX}"); return 1
    print(f"OPEN {INPUT_DOCX}", flush=True)
    doc = Document(INPUT_DOCX)

    # PREFLIGHT: prose targets each present once; Table 5.10 has expected shape
    ok = True
    for old, _new in PROSE_EDITS:
        n = count_occ(doc, old)
        if n != 1:
            ok = False; print(f"  MISMATCH prose count={n}: '{old[:45]}...'", flush=True)
    t = doc.tables[TABLE_510_IDX]
    if t.rows[0].cells[0].text.strip() != "Query" or len(t.rows) < 7:
        ok = False; print(f"  MISMATCH Table {TABLE_510_IDX} shape: hdr='{t.rows[0].cells[0].text}' rows={len(t.rows)}", flush=True)
    if not ok:
        print("ABORT preflight mismatch — NOT saving.", flush=True); return 2
    print("PREFLIGHT ok", flush=True)

    # PHASE 1: Table 5.10 rebuild
    for ri, q, g, f in TABLE_510_ROWS:
        set_cell(t.rows[ri].cells[0], q)
        set_cell(t.rows[ri].cells[1], g)
        set_cell(t.rows[ri].cells[2], f)
    print(f"TABLE 5.10 rebuilt ({len(TABLE_510_ROWS)} rows)", flush=True)

    # PHASE 2: prose
    paras = list(all_paras(doc))
    for old, new in PROSE_EDITS:
        tot = sum(replace_in_paragraph(p, old, new) for p in paras)
        if tot != 1:
            print(f"ABORT prose '{old[:30]}' applied={tot}; NOT saving.", flush=True); return 3
    print(f"PROSE {len(PROSE_EDITS)} edits applied", flush=True)

    # POSTFLIGHT
    for tok in ("How do the courts contrast the dominancy test with the holistic test?",
                "role of the Intellectual Property Office of the Philippines?",
                "dominancy and holistic tests for trademark infringement",
                "reproduced grounding at 0.993",
                "subset of the eighteen-query re-evaluation set"):
        if count_occ(doc, tok) == 0:
            print(f"FAIL expected token missing: {tok[:40]}", flush=True); return 4
    for gone in ("What are the key findings?", "Who are the main researchers?",
                 "constitutional rights across decisions", "datasets and concepts"):
        if count_occ(doc, gone) != 0:
            print(f"FAIL stale token still present: {gone}", flush=True); return 5
    # Table 5.10 mean must equal mean of its rows now
    mean_g = t.rows[6].cells[1].text.strip()
    if mean_g != "0.984":
        print(f"FAIL Table 5.10 mean G = {mean_g} (expected 0.984)", flush=True); return 6
    heads = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    for req in ("Chapter 7: Conclusion", "References", "Appendix A: Glossary of Terms",
                "5.8.1 Post-Reroll"):
        if not any(req in h for h in heads):
            print(f"FAIL heading missing: {req}", flush=True); return 7
    print(f"POSTFLIGHT ok | tables={len(doc.tables)} paras={len(doc.paragraphs)}", flush=True)

    doc.save(OUTPUT_DOCX)
    print(f"SAVED {OUTPUT_DOCX}", flush=True)
    print("DONE — Table 5.10 = real IP queries, §3.3.3 honest, mean-bug fixed, 2nd run noted.", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
