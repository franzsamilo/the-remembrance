"""Update v6.3 -> v6.4 with story-hierarchy improvements and new visualizations.

Additions (all additive, nothing removed):
1. Insert Figure 5.0 (architecture diagram) at start of Chapter 5
2. Insert "Executive Summary at a Glance" callout/table at top of Ch5
3. Insert "Chapter 5 Reading Guide" navigation table
4. Insert Figure 5.5 (decision tree) at top of section 5.2
5. Insert Figure 5.6 (hypothesis dashboard) at end of section 5.10
6. Insert Figure 5.7 (uplift evolution) in section 5.9
7. Add Figure 5.0 cross-reference in Chapter 3.2

Output: Project Study Report_ The Remembrance 6.4.docx
"""
from __future__ import annotations

import os
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(PROJECT_ROOT, "Project Study Report_ The Remembrance 6.3.docx")
OUTPUT_DOCX = os.path.join(PROJECT_ROOT, "Project Study Report_ The Remembrance 6.4.docx")
FIGURES_DIR = os.path.join(PROJECT_ROOT, "docs", "paper", "figures")


def _set_table_borders(table, color="000000", size="4"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tblBorders.append(border)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _set_cell_shading(cell, color_hex):
    """Set cell background color (e.g., '2D6A4F' for green)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _add_paragraph_before(reference_paragraph, doc, text="", style=None):
    new_p = doc.add_paragraph(text, style=style)
    reference_paragraph._p.addprevious(new_p._p)
    return new_p


def _add_table_before(reference_paragraph, doc, rows, cols, border_color="000000", border_size="4"):
    new_t = doc.add_table(rows=rows, cols=cols)
    _set_table_borders(new_t, color=border_color, size=border_size)
    reference_paragraph._p.addprevious(new_t._tbl)
    return new_t


def _add_picture_before(reference_paragraph, doc, image_path, width_inches=6.5):
    new_p = doc.add_paragraph()
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    reference_paragraph._p.addprevious(new_p._p)
    return new_p


def _set_cell_bold(cell, bold=True):
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = bold


def _fill_table(table, header, rows_data, bold_header=True, bold_columns=None):
    bold_columns = bold_columns or []
    for j, h in enumerate(header):
        c = table.rows[0].cells[j]
        c.text = h
        if bold_header:
            _set_cell_bold(c, True)
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            c = table.rows[i + 1].cells[j]
            c.text = str(val)
            if j in bold_columns:
                _set_cell_bold(c, True)


def _find_paragraph_by_text(doc, text_substr):
    for p in doc.paragraphs:
        if text_substr in p.text:
            return p
    return None


def main():
    doc = Document(INPUT_DOCX)

    # ============================================================
    # Find anchor paragraphs
    # ============================================================
    ch5_heading = None
    for p in doc.paragraphs:
        if (p.style and p.style.name == "Heading 1"
                and p.text.strip().startswith("Chapter 5:")):
            ch5_heading = p
            break
    if ch5_heading is None:
        raise RuntimeError("Could not find Chapter 5 heading")

    ch5_intro = None
    # Find the first non-empty paragraph after Chapter 5 heading.
    # NOTE: doc.paragraphs returns a fresh list each call; identity check
    # fails across iterations. Compare underlying XML element instead.
    paras_list = list(doc.paragraphs)
    found_ch5_idx = None
    for idx, p in enumerate(paras_list):
        if p._p is ch5_heading._p:
            found_ch5_idx = idx
            break
    if found_ch5_idx is not None:
        for p in paras_list[found_ch5_idx + 1:]:
            if p.text.strip():
                ch5_intro = p
                break
    if ch5_intro is None:
        raise RuntimeError("Could not find Chapter 5 intro paragraph")

    sec_5_1_heading = _find_paragraph_by_text(doc, "5.1 Corpus Statistics")
    sec_5_2_heading = _find_paragraph_by_text(doc, "5.2 Tuning Campaign Overview")
    sec_5_2_intro = None
    if sec_5_2_heading is not None:
        paras_list2 = list(doc.paragraphs)
        for idx, p in enumerate(paras_list2):
            if p._p is sec_5_2_heading._p:
                for q in paras_list2[idx + 1:]:
                    if q.text.strip():
                        sec_5_2_intro = q
                        break
                break

    sec_5_9_intro = _find_paragraph_by_text(
        doc, "To test H1 (Topological Correlation")

    sec_6_heading = None
    for p in doc.paragraphs:
        if (p.style and p.style.name == "Heading 1"
                and p.text.strip().startswith("Chapter 6:")):
            sec_6_heading = p
            break

    sec_3_2_2_arch = _find_paragraph_by_text(
        doc, "Initial baseline (Run 1) used a 2-layer")

    # ============================================================
    # 1. INSERT Figure 5.0 (Architecture) at start of Chapter 5
    # ============================================================
    # Insert AFTER ch5_intro (so the intro paragraph stays first)
    # We insert before sec_5_1_heading to put it between intro and 5.1
    if sec_5_1_heading is not None:
        _add_paragraph_before(sec_5_1_heading, doc, "Figure 5.0: System Architecture Overview")
        _add_picture_before(sec_5_1_heading, doc,
                            os.path.join(FIGURES_DIR, "fig5_0_architecture.png"),
                            width_inches=6.8)
        _add_paragraph_before(sec_5_1_heading, doc,
            "Figure 5.0 visualizes the three-pipeline 'Validate-then-Generate' architecture "
            "described in Chapter 3. The Feature Pipeline (top, green) ingests PDFs through "
            "SimpleKGPipeline, extracts entities and relationships via Gemini 2.5 Flash, stores "
            "them in Neo4j, and embeds nodes with DistilBERT. The Training Pipeline (middle, gold) "
            "runs the CompGCN integrity audit and writes plausibility scores to all 6,419 "
            "non-FROM_CHUNK edges. The Inference Pipeline (bottom, red) performs hybrid retrieval, "
            "passes the retrieved subgraph through the GNN filter (highlighted), and synthesizes a "
            "grounded narrative or returns a hard 'Grounding Error' when no triplets pass the "
            "threshold. The boldface 'GNN FILTER' box is the architectural contribution of this "
            "work — it is the gate that separates retrieval from synthesis."
        )
        print("[1] Figure 5.0 (architecture diagram) inserted at start of Chapter 5")

    # ============================================================
    # 2. INSERT "Executive Summary at a Glance" table at top of Ch5
    #    (Inserted between Ch5 heading and the existing intro paragraph)
    # ============================================================
    summary_heading = _add_paragraph_before(ch5_intro, doc,
        "Executive Summary: Headline Results", style="Heading 2")
    _add_paragraph_before(ch5_intro, doc,
        "The empirical evaluation reported in this chapter culminates in the Run 8 recommended "
        "configuration achieving all four paper KPIs at the canonical generator-side filter "
        "threshold (tau = 0.95). The headline numbers are presented below; detailed experimental "
        "narrative and methodology follow in Sections 5.1 through 5.10."
    )

    summary_table = _add_table_before(ch5_intro, doc, rows=6, cols=4,
                                       border_color="2D6A4F", border_size="8")
    _fill_table(summary_table,
        ["Hypothesis / KPI", "Paper Target", "Achieved (Run 8)", "Status"],
        [
            ["H1: Topological Correlation",
             "Detectable GNN uplift",
             "+45% Grounding, +204% Faithfulness vs prompt-only",
             "CONFIRMED"],
            ["H2: AUC-ROC",
             "> 0.95",
             "0.985 +/- 0.001 (12-seed mean +/- std)",
             "PASS"],
            ["H2: MRR (uniform eval)",
             "> 0.95",
             "0.958 +/- 0.005 (12-seed mean +/- std)",
             "PASS"],
            ["H3: Grounding Score",
             "> 0.98",
             "0.9884 (tau = 0.95)",
             "PASS"],
            ["H3: Faithfulness",
             "high (paper specification)",
             "0.9714 (tau = 0.95)",
             "PASS"],
        ], bold_columns=[2, 3])

    # Color the status column green for visual emphasis
    for i in range(1, 6):
        _set_cell_shading(summary_table.rows[i].cells[3], "E6F4EA")

    _add_paragraph_before(ch5_intro, doc,
        "Recommended thesis-defense configuration: 3-layer CompGCN encoder with LayerNorm "
        "(introduced in Run 2); BPR loss with self-adversarial negative weighting at "
        "alpha = 1.0 (introduced in Run 8); DistMult decoder (confirmed via decoder ablation in "
        "Run 9); uniform random negative sampling (default; type-aware tested in Run 7 and "
        "reverted); generator-side filtering at tau = 0.95. Multi-seed evaluation methodology "
        "(n = 12) used for AUC-ROC and MRR — see Section 5.3 and Section 6.3."
    )

    # Reading guide table
    _add_paragraph_before(ch5_intro, doc, "Chapter 5 Reading Guide", style="Heading 3")
    _add_paragraph_before(ch5_intro, doc,
        "Readers seeking the quickest path to the headline results should consult Sections 5.2, "
        "5.3, and 5.10. Readers interested in the experimental ablations and design rationale "
        "should read sequentially. Sections 5.6 (decoder ablation, RotatE regression) and 5.7 "
        "(loss-dependent threshold calibration) are essential context for the discussion in "
        "Chapter 6."
    )

    guide_table = _add_table_before(ch5_intro, doc, rows=6, cols=3)
    _fill_table(guide_table,
        ["Reader Goal", "Read These Sections", "Key Artifact"],
        [
            ["Headline result only",
             "5.2, 5.3, 5.10",
             "Table 5.12, Figure 5.6"],
            ["Why Run 8 wins",
             "5.4, 5.5, 5.6",
             "Tables 5.5, 5.6, 5.8"],
            ["Why we trust the numbers",
             "5.3, 5.7",
             "Table 5.4, Figure 5.2"],
            ["Architectural contribution validation",
             "5.8, 5.9",
             "Tables 5.10, 5.11, Figure 5.4"],
            ["Limitations and future work",
             "Chapter 6 (especially 6.5)",
             "5-item priority list"],
        ])

    print("[2] Executive Summary + Reading Guide inserted at top of Chapter 5")

    # ============================================================
    # 3. INSERT Figure 5.5 (Decision Tree) at start of section 5.2
    # ============================================================
    if sec_5_2_intro is not None:
        # Insert after sec_5_2_intro paragraph (i.e., between the intro and Table 5.3)
        # Better: insert at start of section 5.2 — between heading and intro
        _add_paragraph_before(sec_5_2_intro, doc, "Figure 5.5: Tuning Campaign Decision Tree")
        _add_picture_before(sec_5_2_intro, doc,
                            os.path.join(FIGURES_DIR, "fig5_5_decision_tree.png"),
                            width_inches=6.8)
        _add_paragraph_before(sec_5_2_intro, doc,
            "Figure 5.5 visualizes the campaign as a decision tree. Each box represents a single "
            "tuning run; green-bordered boxes indicate interventions that were retained in the final "
            "configuration; red-bordered boxes indicate interventions that were tested and reverted "
            "or regressed; gray-bordered boxes indicate infrastructure improvements that did not "
            "themselves change the model. The tree converges on Run 8 as the recommended "
            "configuration."
        )
        print("[3] Figure 5.5 (decision tree) inserted at start of section 5.2")

    # ============================================================
    # 4. INSERT improved Figure 5.3 v2 in section 5.4
    # ============================================================
    fig_5_3_caption = _find_paragraph_by_text(
        doc, "Figure 5.3: Score Distribution by Loss")
    if fig_5_3_caption is not None:
        # Add the v2 figure right after the existing one
        # Find next paragraph after this caption (which should be the embedded image),
        # then insert v2 caption + image before the NEXT description paragraph
        # Easier: insert before the v3 explanation paragraph
        v3_explain = _find_paragraph_by_text(
            doc, "Figure 5.3 visualizes the count of edges falling")
        if v3_explain is not None:
            # Skip past the explanation; find the next section (5.5)
            sec_5_5_heading = _find_paragraph_by_text(
                doc, "5.5 Negative Sampling Ablation")
            if sec_5_5_heading is not None:
                _add_paragraph_before(sec_5_5_heading, doc,
                    "Figure 5.3b: Score Distribution (Proportional View + Log Scale)")
                _add_picture_before(sec_5_5_heading, doc,
                                    os.path.join(FIGURES_DIR, "fig5_3_score_distribution_v2.png"),
                                    width_inches=6.8)
                _add_paragraph_before(sec_5_5_heading, doc,
                    "Figure 5.3b presents two complementary views of the same data. The left panel "
                    "shows the proportional distribution of edges by score bucket as a stacked bar "
                    "chart, with non-trivial percentages annotated. The right panel uses a "
                    "logarithmic y-axis to make Run 9's RotatE-induced score-range collapse "
                    "visible alongside the other configurations - all 6,419 edges fall below 0.50, "
                    "rendering the canonical paper threshold (red shaded region) effectively "
                    "unreachable. This visualization motivates the loss-dependent calibration "
                    "discussion in Section 5.7 and the architectural-robustness observation in "
                    "Section 6.4."
                )
                print("[4] Figure 5.3b (improved score distribution) inserted before section 5.5")

    # ============================================================
    # 5. INSERT Figure 5.7 (Uplift Evolution) in section 5.9
    # ============================================================
    fig_5_4_caption = _find_paragraph_by_text(
        doc, "Figure 5.4: GNN Uplift over Prompt-Only Baseline")
    if fig_5_4_caption is not None:
        # Insert Figure 5.7 right after the Figure 5.4 explanation
        sec_5_10_heading = _find_paragraph_by_text(doc, "5.10 Hypothesis Validation Summary")
        if sec_5_10_heading is not None:
            _add_paragraph_before(sec_5_10_heading, doc,
                "Figure 5.7: GNN Uplift Evolution Across the Campaign")
            _add_picture_before(sec_5_10_heading, doc,
                                os.path.join(FIGURES_DIR, "fig5_7_uplift_evolution.png"),
                                width_inches=6.5)
            _add_paragraph_before(sec_5_10_heading, doc,
                "Figure 5.7 plots Grounding (green) and Faithfulness (blue) uplift percentages "
                "over the prompt-only baseline across three representative runs. The uplift grows "
                "monotonically as the model improves, from +10%/+26% at Run 1 (baseline) to "
                "+45%/+204% at Run 8 (recommended). This monotone growth is the cleanest "
                "empirical signal supporting H1 (Topological Correlation) - better learned "
                "topology produces a larger generative-quality benefit, not merely a better "
                "internal scoring metric."
            )
            print("[5] Figure 5.7 (uplift evolution) inserted before section 5.10")

    # ============================================================
    # 6. INSERT Figure 5.6 (Hypothesis Dashboard) at end of section 5.10
    #    (right before Chapter 6 heading)
    # ============================================================
    if sec_6_heading is not None:
        _add_paragraph_before(sec_6_heading, doc,
            "Figure 5.6: Hypothesis Validation Dashboard")
        _add_picture_before(sec_6_heading, doc,
                            os.path.join(FIGURES_DIR, "fig5_6_hypothesis_dashboard.png"),
                            width_inches=6.8)
        _add_paragraph_before(sec_6_heading, doc,
            "Figure 5.6 presents the final hypothesis validation status as a dashboard. All four "
            "numeric paper KPIs (H2 AUC-ROC, H2 MRR, H3 Grounding, H3 Faithfulness) clear their "
            "respective targets at the canonical generator-side filter threshold of tau = 0.95. "
            "H1 (Topological Correlation) is qualitatively confirmed by the GNN uplift over "
            "prompt-only baseline reported in Section 5.9. The empirical evidence supports the "
            "central architectural claim of this thesis: a CompGCN-based integrity layer "
            "interposed between retrieval and synthesis materially reduces hallucination while "
            "preserving retrieval coverage and providing an explicit Grounding Error refusal "
            "signal when validated triplets are unavailable."
        )
        print("[6] Figure 5.6 (hypothesis dashboard) inserted before Chapter 6")

    # ============================================================
    # 7. ADD Figure 5.0 cross-reference in Chapter 3.2
    # ============================================================
    if sec_3_2_2_arch is not None:
        # Append a cross-reference line to the architecture description paragraph
        original = sec_3_2_2_arch.text
        sec_3_2_2_arch.text = original.rstrip() + (
            " A consolidated visualization of the three-pipeline architecture is presented as "
            "Figure 5.0 in Chapter 5."
        )
        print("[7] Cross-reference to Figure 5.0 added in Chapter 3.2.2")

    # ============================================================
    # SAVE
    # ============================================================
    doc.save(OUTPUT_DOCX)
    print(f"\nSaved: {OUTPUT_DOCX}")
    print(f"Final: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")


if __name__ == "__main__":
    main()
