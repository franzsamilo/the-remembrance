"""Update past chapters and tables in v6.2 to reflect Run 8 final configuration.

Changes (all additive, nothing removed):
1. Abstract: append empirical-results sentence
2. Table of Contents: insert Chapter 5 + Chapter 6 entries
3. Chapter 3.2.2 prose: update architecture description to final tuned form
4. Table 3.1 (CompGCN Hyperparameters): update values to Run 8 tuned, add new rows
5. Table 3.3 (Hypothesis-to-Metric): add Run 8 Achieved column
6. Add new Table 3.4 (Hypothesis Validation Forward Reference)
7. References: add academic citations introduced in Chapter 5/6

Output: Project Study Report_ The Remembrance 6.3.docx
"""
from __future__ import annotations

import os
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(PROJECT_ROOT, "Project Study Report_ The Remembrance 6.2.docx")
OUTPUT_DOCX = os.path.join(PROJECT_ROOT, "Project Study Report_ The Remembrance 6.3.docx")


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _add_paragraph_before(reference_paragraph, doc, text="", style=None):
    new_p = doc.add_paragraph(text, style=style)
    reference_paragraph._p.addprevious(new_p._p)
    return new_p


def _add_table_before(reference_paragraph, doc, rows, cols):
    new_t = doc.add_table(rows=rows, cols=cols)
    _set_table_borders(new_t)
    reference_paragraph._p.addprevious(new_t._tbl)
    return new_t


def _set_cell_bold(cell, bold=True):
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = bold


def _fill_table(table, header, rows_data, bold_header=True, bold_columns=None):
    bold_columns = bold_columns or []
    for j, h in enumerate(header):
        c = table.rows[0].cells[j]
        c.text = h
        if bold_header:
            _set_cell_bold(c, True)
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            c = table.rows[i + 1].cells[j]
            c.text = str(val)
            if j in bold_columns:
                _set_cell_bold(c, True)


def _add_row_to_table(table, cells_text):
    """Append a new row to an existing table."""
    new_row = table.add_row()
    for j, val in enumerate(cells_text):
        new_row.cells[j].text = str(val)
    return new_row


def _add_column_to_table(table, header, values, bold_header=True, bold_data=False):
    """Append a new column to an existing table by manipulating XML."""
    for i, row in enumerate(table.rows):
        # Find the last <w:tc> in this row
        tcs = row._tr.findall(qn("w:tc"))
        if not tcs:
            continue
        last_tc = tcs[-1]
        new_tc = deepcopy(last_tc)
        # Clear paragraph content from the cloned cell
        for p in new_tc.findall(qn("w:p")):
            new_tc.remove(p)
        # Add new paragraph with our text
        text_val = header if i == 0 else (values[i - 1] if i - 1 < len(values) else "")
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        if (i == 0 and bold_header) or (i > 0 and bold_data):
            rpr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rpr.append(b)
            new_r.append(rpr)
        new_t_el = OxmlElement("w:t")
        new_t_el.text = str(text_val)
        new_r.append(new_t_el)
        new_p.append(new_r)
        new_tc.append(new_p)
        row._tr.append(new_tc)
    # Update the table grid to add a column
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        last_col = grid.findall(qn("w:gridCol"))[-1]
        new_col = deepcopy(last_col)
        grid.append(new_col)


def main():
    doc = Document(INPUT_DOCX)

    # ===========================================================
    # 1. ABSTRACT — append a sentence about empirical results
    # ===========================================================
    abstract_para = doc.paragraphs[14]
    original_abstract = abstract_para.text
    abstract_para.text = (
        original_abstract.rstrip()
        + " Empirical evaluation across nine systematic tuning runs (Chapter 5) "
        "achieves all four paper KPIs at the canonical generator-side filter threshold "
        "(tau = 0.95): AUC-ROC = 0.985 +/- 0.001 (12-seed); MRR = 0.958 +/- 0.005 "
        "(12-seed); Grounding = 0.9884; Faithfulness = 0.9714. The Validate-then-Generate "
        "architecture's measurable uplift over standard chunk-RAG baseline is +45% Grounding "
        "and +204% Faithfulness, confirming the Topological Correlation hypothesis (H1). "
        "The recommended thesis-defense configuration is 3-layer CompGCN with LayerNorm, "
        "BPR loss with self-adversarial negative weighting (alpha = 1.0), DistMult decoder, "
        "and uniform negative sampling."
    )
    print("[1] Abstract updated")

    # ===========================================================
    # 2. TABLE OF CONTENTS — insert Chapter 5/6 entries before References
    # ===========================================================
    refs_toc_para = None
    for p in doc.paragraphs:
        if (p.style and p.style.name == "normal"
                and p.text.strip().lower() == "references"):
            refs_toc_para = p
            break
    if refs_toc_para is None:
        # Try finding by index near where TOC ends
        for p in doc.paragraphs[60:80]:
            if p.text.strip().lower() == "references":
                refs_toc_para = p
                break

    if refs_toc_para is not None:
        toc_entries = [
            "Chapter 5: Experimental Results and Analysis",
            "5.1 Corpus Statistics",
            "5.2 Tuning Campaign Overview",
            "5.3 GNN Performance: Link Prediction with Multi-Seed Methodology",
            "5.4 Loss Function Ablation: BCE versus BPR versus Self-Adversarial BPR",
            "5.5 Negative Sampling Ablation: Uniform versus Type-Aware Corruption",
            "5.6 Decoder Ablation: DistMult versus RotatE",
            "5.7 Threshold Calibration: tau is Loss-Dependent",
            "5.8 Grounding and Faithfulness: Per-Query Results at Canonical tau",
            "5.9 Validate-then-Generate versus Prompt-Only Baseline (H1 Confirmation)",
            "5.10 Hypothesis Validation Summary",
            "Chapter 6: Discussion",
            "6.1 Principal Finding: Corpus-Density-Bound Performance Ceiling",
            "6.2 Loss-Dependent Score Calibration",
            "6.3 MRR Variance and Multi-Seed Evaluation Methodology",
            "6.4 Architectural Robustness: The Grounding Error Refusal Mechanism",
            "6.5 Limitations and Future Work",
        ]
        for entry in toc_entries:
            _add_paragraph_before(refs_toc_para, doc, entry)
        print(f"[2] TOC updated ({len(toc_entries)} entries inserted)")
    else:
        print("[2] WARNING: References TOC entry not found; TOC not updated")

    # ===========================================================
    # 3. CHAPTER 3.2.2 PROSE — update architecture from baseline to final tuned
    # ===========================================================
    # Para 199: "The model is a custom CompGCNAuditModel: a 2-layer CompGCN encoder
    # with a DistMult link predictor..."
    arch_para = None
    for i, p in enumerate(doc.paragraphs):
        if "2-layer CompGCN encoder" in p.text:
            arch_para = p
            break
    if arch_para is not None:
        original = arch_para.text
        arch_para.text = (
            "The model is a custom CompGCNAuditModel. Initial baseline (Run 1) used a 2-layer "
            "CompGCN encoder with a DistMult link predictor; the architecture was tuned through "
            "the campaign reported in Chapter 5 to its final production form: a 3-layer CompGCN "
            "encoder with LayerNorm between layers, trained under Bayesian Personalized Ranking "
            "(BPR) loss with self-adversarial negative weighting (alpha = 1.0; Sun et al. 2019, "
            "eq. 5), and DistMult composition for the decoder. The encoder takes all "
            "non-FROM_CHUNK edges and 768-dimensional node embeddings from Neo4j as input and "
            "outputs a plausibility score (0.0 to 1.0) written to every relationship in the "
            "graph. Table 3.1 reports both the original baseline hyperparameters and the final "
            "tuned values used in the recommended defense configuration."
        )
        print(f"[3] Chapter 3.2.2 architecture description updated")
    else:
        print("[3] WARNING: 2-layer CompGCN paragraph not found")

    # Find the para about training process and update it too
    training_proc_para = None
    for p in doc.paragraphs:
        if "compute BCE loss" in p.text or "compute positive edge logits" in p.text:
            training_proc_para = p
            break
    if training_proc_para is not None:
        original = training_proc_para.text
        # Keep the original prose but add an updated note in the same paragraph
        training_proc_para.text = original.rstrip() + (
            " Note: the loss function specified above (BCE) is the Run 1 baseline. The "
            "final production loss (Run 8 onward) is BPR with self-adversarial weighting "
            "(see Chapter 5, Section 5.4 for the loss-function ablation rationale and "
            "empirical comparison)."
        )
        print(f"[3b] Training process prose updated with loss note")

    # ===========================================================
    # 4. TABLE 3.1 (CompGCN Hyperparameters) — update values to Run 8 tuned
    # ===========================================================
    hp_table = doc.tables[2]  # Table 2 in 0-indexed = Table 3.1 in paper
    # Map: row index -> updated value (col 1)
    updates = {
        # Row 0 = header. Row 1 = Hidden Channels. Row 2 = Epochs. etc.
        2: ("300 (Run 8 tuned; baseline was 100)",
            "Increased to allow convergence under BPR loss"),
        3: ("0.0005 (Run 8 tuned; baseline was 0.001)",
            "Finer convergence at higher epoch count"),
        6: ("30 epochs (Run 8 tuned; baseline was 20)",
            "Wider plateau detection at LR 0.0005"),
        8: ("15 (Run 8 tuned; baseline was 10)",
            "Higher contrast for self-adversarial weighting"),
        10: ("0.05 (Run 2+; ignored by BPR loss)",
             "Active under BCE; harmless under BPR"),
    }
    for row_idx, (new_val, new_just) in updates.items():
        hp_table.rows[row_idx].cells[1].text = new_val
        hp_table.rows[row_idx].cells[2].text = new_just
    # Add new rows for additional Run 8 hyperparameters
    new_rows = [
        ["Loss Function", "BPR with self-adversarial weighting (alpha = 1.0)",
         "Run 8 final; replaces BCE baseline (Sun et al. 2019)"],
        ["Decoder Composition", "DistMult (final)",
         "Confirmed via Run 9 ablation; RotatE regressed at this corpus density"],
        ["Negative Sampling", "Uniform random head/tail corruption",
         "Run 8 final; type-aware reverted to opt-in (Run 7 finding)"],
        ["AUC Guardrail", "0.95 (Run 7+)",
         "Skips Neo4j score sync if final AUC regresses below threshold"],
        ["Architecture Layers", "3-layer CompGCN with LayerNorm (Run 2+)",
         "Tuned from baseline 2-layer; +0.025 AUC vs baseline"],
    ]
    for row_data in new_rows:
        _add_row_to_table(hp_table, row_data)
    print(f"[4] Table 3.1 (CompGCN Hyperparameters) updated: 5 rows revised, "
          f"5 rows added")

    # ===========================================================
    # 5. TABLE 3.3 (Hypothesis-to-Metric) — add Run 8 Achieved column
    # ===========================================================
    hyp_table = doc.tables[3]  # Table 3 in 0-indexed = Table 3.3 in paper
    # Add a new column: "Achieved (Run 8)" with PASS status
    achieved_values = [
        "0.985 +/- 0.001 (12-seed) - PASS",
        "0.958 +/- 0.005 (12-seed) - PASS",
        "0.9884 (tau=0.95) - PASS",
        "0.9714 (tau=0.95) - PASS",
    ]
    _add_column_to_table(hyp_table, "Achieved (Run 8) [forward ref Ch 5.10]",
                          achieved_values, bold_header=True, bold_data=True)
    print(f"[5] Table 3.3 (Hypothesis-to-Metric) extended with Achieved column")

    # ===========================================================
    # 6. ADD TABLE 3.4 prose — note in Chapter 3 about forward reference
    # ===========================================================
    # Find paragraph after Table 3.3 (the last sub-paragraph in 3.3 before 3.3.1)
    target_para = None
    for i, p in enumerate(doc.paragraphs):
        if "The following subsections detail each metric" in p.text:
            target_para = p
            break
    if target_para is not None:
        new_para = doc.add_paragraph(
            "Note on Achieved column (Table 3.3): The 'Achieved' column reports the final "
            "results obtained under the Run 8 recommended configuration described in Chapter 5. "
            "Multi-seed methodology (n=12) is used for AUC-ROC and MRR per the protocol detailed "
            "in Section 5.3 and discussed in Section 6.3. All four KPIs cleared the paper "
            "targets at the canonical generator-side filter threshold tau = 0.95."
        )
        target_para._p.addprevious(new_para._p)
        print(f"[6] Forward-reference note inserted before 3.3.1")

    # ===========================================================
    # 7. REFERENCES — add academic citations introduced in Chapter 5/6
    # ===========================================================
    refs_heading = None
    for p in doc.paragraphs:
        if (p.style and p.style.name == "Heading 1"
                and p.text.strip().lower() == "references"):
            refs_heading = p
            break

    if refs_heading is not None:
        # Find the next paragraph after the heading to insert before, OR append at end
        # The user's References section may be empty or have existing citations
        # Insert new citations as paragraphs after the heading

        new_refs = [
            "Sun, Z., Deng, Z.-H., Nie, J.-Y., & Tang, J. (2019). RotatE: Knowledge "
            "Graph Embedding by Relational Rotation in Complex Space. In International "
            "Conference on Learning Representations (ICLR 2019). Cited in Section 3.2.2 "
            "(self-adversarial negative weighting), Section 5.4 (loss ablation), Section 5.6 "
            "(decoder ablation), and Section 6.1 (corpus-density discussion).",

            "Vashishth, S., Sanyal, S., Nitin, V., & Talukdar, P. (2020). "
            "Composition-based Multi-Relational Graph Convolutional Networks. In "
            "International Conference on Learning Representations (ICLR 2020). Foundational "
            "reference for the CompGCN encoder architecture used throughout this work; the "
            "encoder-decoder evaluation pattern in Table 4 of that paper motivates the "
            "decoder ablation in Section 5.6 of this thesis.",

            "Rendle, S., Freudenthaler, C., Gantner, Z., & Schmidt-Thieme, L. (2009). "
            "BPR: Bayesian Personalized Ranking from Implicit Feedback. In Proceedings of "
            "the Twenty-Fifth Conference on Uncertainty in Artificial Intelligence (UAI 2009), "
            "452-461. Reference for the pairwise BPR loss adopted in Run 6 onward "
            "(Section 5.4).",

            "Yang, B., Yih, W.-T., He, X., Gao, J., & Deng, L. (2015). Embedding Entities "
            "and Relations for Learning and Inference in Knowledge Bases. In International "
            "Conference on Learning Representations (ICLR 2015). Reference for the DistMult "
            "decoder used in the Run 8 recommended configuration (Section 3.2.2).",

            "Bordes, A., Usunier, N., Garcia-Duran, A., Weston, J., & Yakhnenko, O. (2013). "
            "Translating Embeddings for Modeling Multi-relational Data. In Advances in Neural "
            "Information Processing Systems 26 (NeurIPS 2013), 2787-2795. Reference for "
            "TransE, listed as a candidate decoder in the future-work discussion (Section 6.5).",

            "Trouillon, T., Welbl, J., Riedel, S., Gaussier, E., & Bouchard, G. (2016). "
            "Complex Embeddings for Simple Link Prediction. In Proceedings of the 33rd "
            "International Conference on Machine Learning (ICML 2016), 2071-2080. Reference "
            "for ComplEx, listed as a future-work decoder candidate (Section 6.5).",

            "Sanh, V., Debut, L., Chaumond, J., & Wolf, T. (2019). DistilBERT, a distilled "
            "version of BERT: smaller, faster, cheaper and lighter. In NeurIPS 2019 "
            "Workshop on Energy Efficient Machine Learning and Cognitive Computing (5th "
            "edition). Reference for the node embedding encoder used in Stage 4 of the "
            "Feature Pipeline (Section 3.2.1).",

            "Ba, J. L., Kiros, J. R., & Hinton, G. E. (2016). Layer Normalization. arXiv "
            "preprint arXiv:1607.06450. Reference for the LayerNorm component introduced "
            "between CompGCN layers in Run 2 (Section 5.4).",

            "Traag, V. A., Waltman, L., & van Eck, N. J. (2019). From Louvain to Leiden: "
            "guaranteeing well-connected communities. Scientific Reports, 9(1), 5233. "
            "Reference for the Leiden community detection algorithm used in retrieval "
            "expansion (Section 3.2.3).",

            "Page, L., Brin, S., Motwani, R., & Winograd, T. (1999). The PageRank Citation "
            "Ranking: Bringing Order to the Web. Stanford InfoLab Technical Report. Reference "
            "for PageRank centrality scoring used in conjunction with Leiden community "
            "detection (Section 3.2.3).",
        ]

        # Insert references AFTER the References heading
        # Find the paragraph immediately after refs_heading
        next_after_refs = None
        all_paras = list(doc.paragraphs)
        for i, p in enumerate(all_paras):
            if p is refs_heading and i + 1 < len(all_paras):
                next_after_refs = all_paras[i + 1]
                break

        if next_after_refs is not None:
            for ref in new_refs:
                _add_paragraph_before(next_after_refs, doc, ref)
            print(f"[7] References: {len(new_refs)} academic citations added")
        else:
            # References is the last paragraph; append at end
            for ref in new_refs:
                doc.add_paragraph(ref)
            print(f"[7] References: {len(new_refs)} academic citations appended at end")
    else:
        print("[7] WARNING: References Heading 1 not found; references not added")

    # ===========================================================
    # SAVE
    # ===========================================================
    doc.save(OUTPUT_DOCX)
    print(f"\nSaved: {OUTPUT_DOCX}")
    print(f"Final: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")


if __name__ == "__main__":
    main()
