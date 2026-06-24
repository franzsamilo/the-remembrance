"""Update v6.9 -> v6.10: re-ground the paper's corpus stats on the live (reproducible)
graph after the Jun-21 graph-loss + reroll, with an HONEST relation-type reframe.

The lost graph (5,187 GNN nodes / 6,419 edges / density 1.24 / 7 distinct relation
types) was rebuilt by re-ingesting the SAME 14-PDF corpus (train13 checkpoint, epoch
228). The live graph is a different *realization*:

  topology     5,825 nodes / 7,426 edges / density 1.27  (~31x sparser than FB15k)
  sub-counts   embedded 5,533 / unlabeled 239
  labels       Concept 3,296 (57%), Entity 1,456 (25%), Method 270 (5%), ...
  relations    8 distinct in the GNN graph = 6 semantic (USES, CONTRADICTS, EXTENDS,
               PROPOSES, EVALUATES, ACHIEVES) + 2 structural (FROM_DOCUMENT,
               NEXT_CHUNK); FROM_CHUNK is excluded by the loader.

Headline GNN KPIs hold (MRR 0.959, AUC 0.986 — within the printed +/-0.005), so the
corpus-density-bound argument is unchanged and mildly reinforced.

DECISION (user): "Honest reframe, keep model" — present 6 semantic + structural/
provenance edges; no retraining; KPIs stand.

HELD (NOT changed here): MRR/AUC numbers (within live noise) and Grounding/Faithfulness
(0.9884/0.9714 describe the lost graph; live 18-query re-measurement is 0.964/0.918,
awaiting an H3 framing decision).

Percentages use count / TOTAL nodes, matching v6.9's actual method (old Concept
2,804/5,187 = 54%, not /labeled).

Output: Project Study Report_ The Remembrance 6.10.docx
"""
from __future__ import annotations

import copy
import os
from docx import Document

PROJECT_ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(PROJECT_ROOT, "Project Study Report_ The Remembrance 6.9.docx")
OUTPUT_DOCX = os.path.join(PROJECT_ROOT, "Project Study Report_ The Remembrance 6.10.docx")

# ---- PHASE 1: topology + derived figures (global, run-aware; tokens are unique) ----
GLOBAL_REPLACEMENTS: list[tuple[str, str, int]] = [
    ("5,187", "5,825", 2),
    ("6,419", "7,426", 4),
    ("1.24", "1.27", 7),
    ("5,000 entities", "5,800 entities", 1),
    ("6,000 relationships", "7,400 relationships", 1),
    ("32 times sparser", "31 times sparser", 2),
]

# ---- PHASE 4/5: prose phrase replaces (unique in context) ----
PROSE_REPLACEMENTS: list[tuple[str, str, int]] = [
    ("54% of labeled nodes", "57% of labeled nodes", 1),
    ("54% Concept dominance", "57% Concept dominance", 1),
    (
        "7 relationship types (USES, CONTRADICTS, EXTENDS, PROPOSES, EVALUATES, ACHIEVES, FROM_CHUNK)",
        "6 semantic relationship types (USES, CONTRADICTS, EXTENDS, PROPOSES, EVALUATES, ACHIEVES) "
        "plus structural/provenance edges (FROM_CHUNK, FROM_DOCUMENT, NEXT_CHUNK)",
        1,
    ),
    ("seven relationship types", "six semantic relationship types", 1),
]

# ---- PHASE 2/3: targeted table cell edits (numbers are NOT unique -> by coordinate) ----
# (table_index, row, col, old_expected, new_value)
CELL_EDITS: list[tuple[int, int, int, str, str]] = [
    # Table 4 (docx idx 4) — Corpus statistics
    (4, 2, 1, "4,902", "5,533"),                       # embedded nodes
    (4, 3, 1, "232", "239"),                           # unlabeled nodes
    (4, 5, 1, "7", "8 (6 semantic + 2 structural)"),   # distinct relation types
    (4, 8, 1, "54%", "57%"),                           # max label dominance (Concept)
    # Table 10 (docx idx 10) — Label distribution
    (10, 1, 1, "2,804", "3,296"), (10, 1, 2, "54%", "57%"),   # Concept
    (10, 2, 1, "1,410", "1,456"), (10, 2, 2, "27%", "25%"),   # Entity
    (10, 3, 1, "217", "270"),     (10, 3, 2, "4%", "5%"),     # Method
    (10, 4, 1, "208", "245"),                                  # Researcher (% stays 4%)
    (10, 5, 1, "203", "223"),                                  # Result    (% stays 4%)
    (10, 6, 1, "53", "40"),                                    # Metric    (% stays 1%)
    (10, 7, 1, "50", "46"),                                    # Dataset   (% stays 1%)
]

# ---- PHASE 6: add the 2 structural relationship rows to the schema table (idx 1) ----
SCHEMA_TABLE_IDX = 1
STRUCTURAL_ROWS = [
    ("Relationship", "FROM_DOCUMENT", "Provenance link to source document"),
    ("Relationship", "NEXT_CHUNK", "Sequential link between adjacent text chunks"),
]

MUST_SURVIVE = ["0.9884", "0.9714", "0.958", "0.985"]  # deferred numbers must remain


def replace_in_paragraph(par, old: str, new: str) -> int:
    count = 0
    while True:
        runs = par.runs
        full = "".join(r.text for r in runs)
        idx = full.find(old)
        if idx == -1:
            return count
        end = idx + len(old)
        pos = 0
        ranges = []
        for r in runs:
            ranges.append((pos, pos + len(r.text), r))
            pos += len(r.text)
        inserted = False
        for rs, re_, r in ranges:
            if re_ <= idx or rs >= end:
                continue
            ls, le = max(idx, rs) - rs, min(end, re_) - rs
            seg = r.text
            r.text = (seg[:ls] + new + seg[le:]) if not inserted else (seg[:ls] + seg[le:])
            inserted = True
        count += 1


def iter_all_paragraphs(doc):
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def count_occurrences(doc, needle: str) -> int:
    return sum(p.text.count(needle) for p in iter_all_paragraphs(doc))


def set_cell_text(cell, new_text: str) -> None:
    """Replace a cell's text, preserving the first run's formatting."""
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def main() -> int:
    if not os.path.exists(INPUT_DOCX):
        print(f"FATAL input not found: {INPUT_DOCX}", flush=True)
        return 1
    print(f"OPEN {INPUT_DOCX}", flush=True)
    doc = Document(INPUT_DOCX)

    # ---------- PREFLIGHT ----------
    print("PREFLIGHT global tokens:", flush=True)
    ok = True
    for old, _n, exp in GLOBAL_REPLACEMENTS + PROSE_REPLACEMENTS:
        got = count_occurrences(doc, old)
        if got != exp:
            ok = False
            print(f"  MISMATCH '{old[:40]}' found={got} expected={exp}", flush=True)
        else:
            print(f"  OK '{old[:40]}' found={got}", flush=True)
    print("PREFLIGHT cell coordinates:", flush=True)
    for ti, ri, ci, old_exp, _new in CELL_EDITS:
        actual = doc.tables[ti].rows[ri].cells[ci].text.strip()
        if actual != old_exp:
            ok = False
            print(f"  MISMATCH T{ti} r{ri}c{ci}: '{actual}' != expected '{old_exp}'", flush=True)
    if not ok:
        print("ABORT preflight mismatch — NOT saving.", flush=True)
        return 2
    print("  all cell coordinates verified", flush=True)

    # ---------- PHASE 1 + 4/5: paragraph/cell text replacements ----------
    paras = list(iter_all_paragraphs(doc))
    for old, new, exp in GLOBAL_REPLACEMENTS + PROSE_REPLACEMENTS:
        total = sum(replace_in_paragraph(p, old, new) for p in paras)
        flag = "OK" if total == exp else "WARN"
        print(f"  REPL {flag} '{old[:38]}' -> applied={total}/{exp}", flush=True)
        if total != exp:
            print("ABORT replacement diverged — NOT saving.", flush=True)
            return 3

    # ---------- PHASE 2/3: targeted cell edits ----------
    for ti, ri, ci, old_exp, new in CELL_EDITS:
        cell = doc.tables[ti].rows[ri].cells[ci]
        set_cell_text(cell, new)
        if cell.text.strip() != new:
            print(f"FAIL cell T{ti} r{ri}c{ci} did not take new value", flush=True)
            return 4
    print(f"  CELLS {len(CELL_EDITS)} edits applied + verified", flush=True)

    # ---------- PHASE 6: add structural rows to schema table ----------
    schema = doc.tables[SCHEMA_TABLE_IDX]
    rows_before = len(schema.rows)
    for vals in STRUCTURAL_ROWS:
        new_tr = copy.deepcopy(schema.rows[15]._tr)        # clone FROM_CHUNK row formatting
        schema.rows[-1]._tr.addnext(new_tr)                # append after current last row
        nr = schema.rows[-1]
        for cidx, v in enumerate(vals):
            set_cell_text(nr.cells[cidx], v)
    print(f"  SCHEMA rows {rows_before} -> {len(schema.rows)} (+{len(STRUCTURAL_ROWS)})", flush=True)

    # ---------- POSTFLIGHT ----------
    print("POSTFLIGHT:", flush=True)
    for tok in ("5,187", "6,419", "1.24", "54%"):
        rem = count_occurrences(doc, tok)
        print(f"  old '{tok}' remaining={rem}", flush=True)
        if rem != 0:
            print(f"FAIL '{tok}' still present.", flush=True)
            return 5
    for keep in MUST_SURVIVE:
        if count_occurrences(doc, keep) == 0:
            print(f"FAIL guard token '{keep}' destroyed.", flush=True)
            return 6
    # spot-check key new values
    for tok in ("5,825", "7,426", "1.27", "5,533", "3,296", "57%", "FROM_DOCUMENT", "NEXT_CHUNK",
                "six semantic relationship types"):
        if count_occurrences(doc, tok) == 0:
            print(f"FAIL expected new token missing: {tok}", flush=True)
            return 7
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    for req in ("How to Read This Thesis", "Chapter 7: Conclusion", "References",
                "Appendix A: Glossary of Terms"):
        if not any(req in h for h in headings):
            print(f"FAIL heading missing: {req}", flush=True)
            return 8
    img = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
    print(f"  TABLES={len(doc.tables)} IMAGES={img} headings OK", flush=True)

    doc.save(OUTPUT_DOCX)
    print(f"SAVED {OUTPUT_DOCX}", flush=True)
    print("DONE — topology + sub-counts + label dist + relation-type reframe applied; "
          "MRR/AUC + grounding/faithfulness HELD.", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
