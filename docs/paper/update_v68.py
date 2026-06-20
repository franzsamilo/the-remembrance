"""Update v6.7 -> v6.8 with a "How to Read This Thesis" reading-guide
front-matter section.

Inserts one new front-matter section between the Table of Contents and
Chapter 1: Introduction. The guide maps each reader-type to the recommended
section path through the document, plus a short conventions note. Aimed at
readers (panel members, fellow students, curious visitors, reproducers) who
have limited time and want to know where to focus.

The section is "Heading 1" styled to match the chapter / abstract / TOC
level above it, so it appears as a top-level entry in the TOC of the next
recompiled version.

Output: Project Study Report_ The Remembrance 6.8.docx
"""
from __future__ import annotations

import os
from docx import Document
from docx.text.paragraph import Paragraph

PROJECT_ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(
    PROJECT_ROOT,
    "Project Study Report_ The Remembrance 6.7.docx",
)
OUTPUT_DOCX = os.path.join(
    PROJECT_ROOT,
    "Project Study Report_ The Remembrance 6.8.docx",
)


def _find_heading_starting_with(doc, prefix: str) -> Paragraph:
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        if p.text.strip().startswith(prefix):
            return p
    raise LookupError(f"Could not locate anchor heading starting with: {prefix!r}")


def _add_paragraph_before(reference: Paragraph, doc: Document, text: str, style: str | None = None) -> Paragraph:
    new_p = doc.add_paragraph(text, style=style)
    reference._p.addprevious(new_p._p)
    return new_p


def _insert_section_before(reference: Paragraph, doc: Document, heading_text: str, heading_style: str, body_paragraphs: list[str]) -> None:
    _add_paragraph_before(reference, doc, heading_text, style=heading_style)
    for body in body_paragraphs:
        _add_paragraph_before(reference, doc, body)


# ──────────────────────────────────────────────────────────────────────────
# How to Read This Thesis — inserts BEFORE Chapter 1 Introduction.
# Placed after the TOC so it's findable from the TOC AND it reads naturally
# as "before you dive into Chapter 1, here's how to navigate."
# ──────────────────────────────────────────────────────────────────────────

GUIDE_HEADING = "How to Read This Thesis"
GUIDE_BODY = [
    (
        "This thesis serves several reader types — defense panel members "
        "doing formal review, fellow students learning the architecture, "
        "external readers reproducing the experiments, and casual visitors "
        "curious about the contribution. The guide below maps each reader "
        "type to a recommended section path. None of the paths are "
        "exclusive — they are entry points, not boundaries."
    ),
    (
        "If you have five minutes and just want the headline. Read the "
        "Abstract on the previous page; then jump to §5.10 (Hypothesis "
        "Validation Summary) for the four KPI results; then Chapter 7 "
        "(Conclusion) for the architectural contribution stated in one "
        "paragraph. Total reading time: 5–7 minutes."
    ),
    (
        "If you are a panel member doing formal review. Read top-to-bottom "
        "but prioritize: §1.2 (Statement of the Problem) for the framing; "
        "§3.1.1 (Corpus Composition) and §3.3 (Evaluation Framework) for "
        "methodology; Chapter 5 in full for empirical evidence; §6.1 "
        "(Corpus-Density-Bound Ceiling) and §6.6 (Selection Bias) for the "
        "defensive analysis; Chapter 7 (Conclusion). Total: 60–75 minutes "
        "for a careful read."
    ),
    (
        "If you are curious about the architecture but not the experiments. "
        "Read §1.1 (Background), Chapter 2 — especially §2.7 with Figures "
        "2.1, 2.2, and 2.3 — and §3.2 (System Architecture: The "
        "Three-Pipeline Design). The Validate-then-Generate pattern is "
        "fully described in those three locations. Skip Chapter 5; the "
        "tuning campaign details will not change your understanding of the "
        "pattern. Total: 25–30 minutes."
    ),
    (
        "If you want to reproduce the experiments. Start at §6.8 "
        "(Reproducibility Statement) for the replication procedure, then "
        "consult §3.2 for the pipeline definition and §3.3.4 "
        "(Operationalization of Hypotheses) for the falsification criteria. "
        "Chapter 5 documents the configuration; the README.md in the "
        "released codebase covers the boot sequence. Total: 30 minutes "
        "for the paper, plus the replication itself (~2 hours on the "
        "documented hardware floor)."
    ),
    (
        "If you are evaluating the methodology. Read §3.3 (Evaluation "
        "Framework and Metrics) including the four subsections; §6.3 "
        "(MRR Variance and Multi-Seed Evaluation Methodology); §6.5 "
        "(Limitations and Future Work); §6.6 (Selection Bias and "
        "Reproducibility); and §6.7 (Threat Model and Adversarial "
        "Considerations). These sections collectively answer the standard "
        "methodological-review questions. Total: 25–30 minutes."
    ),
    (
        "Conventions used throughout. (1) All headline KPI values reported "
        "outside of single-seed contexts are 12-seed means under canonical "
        "KGE methodology (Sun et al. 2019; Vashishth et al. 2020), and are "
        "annotated with standard deviation where applicable. (2) The "
        "plausibility threshold τ = 0.95 is the deployment-grade default; "
        "where alternative thresholds appear, they are explicitly named "
        "(see §5.7 Threshold Calibration and §3.4 Deployment Strategy). "
        "(3) Figures use color encoding throughout — Feature pipeline "
        "components in green, Training pipeline in gold, Inference "
        "pipeline in red. Black-and-white printing preserves layout but "
        "loses semantic color information. (4) Where the paper documents "
        "a negative result (single-seed MRR shortfall in §5.5; Run 9 "
        "RotatE decoder regression in §5.7), the negative result is "
        "preserved in the publication rather than excised, and the "
        "rationale is discussed in §6.6."
    ),
    (
        "A note on the open-source release. The codebase that produced "
        "these results is released under version control alongside the "
        "thesis. The complete document corpus, the integrity-model "
        "checkpoint, the evaluation queries, and the per-run experimental "
        "logs are all included. The system is corpus-agnostic — any "
        "reader may substitute a different PDF set, re-train the integrity "
        "layer, and re-run the evaluation suite end-to-end to verify or "
        "refute the structural findings reported here."
    ),
]


def main() -> int:
    if not os.path.exists(INPUT_DOCX):
        print(f"FATAL input not found: {INPUT_DOCX}", flush=True)
        return 1

    print(f"OPEN {INPUT_DOCX}", flush=True)
    doc = Document(INPUT_DOCX)
    paragraphs_before = len(doc.paragraphs)

    # Insert "How to Read This Thesis" BEFORE Chapter 1 Introduction
    print("INSERT 'How to Read This Thesis' before Chapter 1", flush=True)
    anchor_ch1 = _find_heading_starting_with(doc, "Chapter 1: Introduction")
    _insert_section_before(
        anchor_ch1, doc,
        heading_text=GUIDE_HEADING,
        heading_style="Heading 1",
        body_paragraphs=GUIDE_BODY,
    )

    paragraphs_after = len(doc.paragraphs)
    delta = paragraphs_after - paragraphs_before
    expected_new = 1 + len(GUIDE_BODY)
    print(f"PARAGRAPHS before={paragraphs_before} after={paragraphs_after} delta={delta} expected={expected_new}", flush=True)
    if delta != expected_new:
        print(
            f"WARN paragraph delta mismatch — expected +{expected_new}, got +{delta}. "
            "Inspect output before publishing.",
            flush=True,
        )

    doc.save(OUTPUT_DOCX)
    print(f"SAVED {OUTPUT_DOCX}", flush=True)

    # Verification
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    if GUIDE_HEADING not in headings:
        print(f"FAIL new heading missing: {GUIDE_HEADING}", flush=True)
        return 2
    print(f"VERIFIED inserted: {GUIDE_HEADING}", flush=True)

    # The order must be: Abstract → Table of Contents → How to Read This Thesis → Chapter 1
    expected_order = ["Abstract", "Table of Contents", GUIDE_HEADING, "Chapter 1: Introduction"]
    actual_seq = []
    for h in headings:
        if h in expected_order and (not actual_seq or h != actual_seq[-1]):
            actual_seq.append(h)
            if len(actual_seq) == len(expected_order):
                break
    if actual_seq != expected_order:
        print(
            f"FAIL front-matter order is wrong.\n  expected: {expected_order}\n  actual:   {actual_seq}",
            flush=True,
        )
        return 3
    print(f"VERIFIED front-matter order: {' -> '.join(expected_order)}", flush=True)

    # Existing v6.7 content still present
    must_preserve = [
        "3.1.1 Corpus Composition and Selection Criteria",
        "3.3.3 Evaluation Query Sample Size Disclosure",
        "3.3.4 Operationalization of Hypotheses",
        "6.6 Selection Bias and Reproducibility",
        "6.7 Threat Model and Adversarial Considerations",
        "6.8 Reproducibility Statement",
        "Chapter 7: Conclusion",
        "References",
    ]
    for req in must_preserve:
        if not any(req in h for h in headings):
            print(f"FAIL existing heading missing after edit: {req}", flush=True)
            return 4
    print("VERIFIED all v6.7 additions preserved", flush=True)

    body_text = ' '.join(p.text for p in doc.paragraphs)
    user_markers = [
        "Figure 2.1: Architectural Comparison",
        "Figure 2.2: Standard RAG Pipeline",
        "Figure 2.3: Validate-then-Generate Pipeline",
        "Figure 3.1: Three-Pipeline Architecture",
        "Table 5.1: Corpus Statistics",
    ]
    missing = [m for m in user_markers if m not in body_text]
    if missing:
        print("FAIL user content from v6.6 absent:", flush=True)
        for m in missing:
            print(f"  - {m}", flush=True)
        return 5
    print("VERIFIED user's Figure 2.x / 3.1 / Table 5.x content preserved", flush=True)

    img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
    table_count = len(doc.tables)
    print(f"EMBEDDED IMAGES: {img_count}", flush=True)
    print(f"TABLES: {table_count}", flush=True)
    print("DONE", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
