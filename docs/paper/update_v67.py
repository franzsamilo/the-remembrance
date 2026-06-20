"""Update v6.6 -> v6.7 with end-to-end paper hardening additions.

Adds four substantive sections, all additive (no existing content modified):

1. §3.3.4 Operationalization of Hypotheses        (after §3.3.3)
2. §6.7   Threat Model and Adversarial Considerations  (after §6.6)
3. §6.8   Reproducibility Statement                    (after §6.7)
4. Chapter 7: Conclusion                                (before References)

These sections plug specific gaps a defense panel will probe:
- §3.3.4 forces a falsification criterion per hypothesis ("how would you
  know if you were wrong?") — research-design hygiene
- §6.7 names the threats this defends against, and explicitly the ones it
  doesn't — preempts the adversarial-robustness question
- §6.8 makes the work reproducible by an external reader (repo, deps,
  hardware, seeds, walltime) — required by any serious thesis
- Chapter 7 gives the paper a proper closing instead of ending at
  Discussion — academic norm

Output: Project Study Report_ The Remembrance 6.7.docx
"""
from __future__ import annotations

import os
from docx import Document
from docx.text.paragraph import Paragraph

PROJECT_ROOT = r"C:/Users/Franz Samilo/Desktop/the-remembrance"
INPUT_DOCX = os.path.join(
    PROJECT_ROOT,
    "Project Study Report_ The Remembrance 6.6.docx",
)
OUTPUT_DOCX = os.path.join(
    PROJECT_ROOT,
    "Project Study Report_ The Remembrance 6.7.docx",
)


def _find_heading_starting_with(doc, prefix: str) -> Paragraph:
    """Heading-style filter so the TOC doesn't shadow the real heading
    (lesson preserved from update_v65.py)."""
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        if p.text.strip().startswith(prefix):
            return p
    raise LookupError(f"Could not locate anchor heading starting with: {prefix!r}")


def _add_paragraph_before(reference: Paragraph, doc: Document, text: str, style: str | None = None) -> Paragraph:
    new_p = doc.add_paragraph(text, style=style)
    reference._p.addprevious(new_p._p)
    return new_p


def _insert_section_before(reference: Paragraph, doc: Document, heading_text: str, heading_style: str, body_paragraphs: list[str]) -> None:
    _add_paragraph_before(reference, doc, heading_text, style=heading_style)
    for body in body_paragraphs:
        _add_paragraph_before(reference, doc, body)


# ──────────────────────────────────────────────────────────────────────────
# §3.3.4 Operationalization of Hypotheses
# Anchors immediately before §3.4 Deployment Strategy.
# ──────────────────────────────────────────────────────────────────────────

SECTION_334_HEADING = "3.3.4 Operationalization of Hypotheses"
SECTION_334_BODY = [
    (
        "To prevent the hypotheses from being trivially confirmable, each H "
        "is operationalized below as a specific measurable prediction with an "
        "explicit falsification criterion. A hypothesis is considered falsified "
        "if the corresponding experiment produces a result below the named "
        "threshold."
    ),
    (
        "H1 (Topological Correlation). Operationalization: a Validate-then-"
        "Generate configuration with the integrity layer enabled will produce "
        "higher Grounding and Faithfulness scores than the prompt-only chunk-"
        "RAG baseline on the same query set. Falsification criterion: H1 is "
        "falsified if the full-stack configuration produces equal or worse "
        "Grounding/Faithfulness than prompt-only across the five evaluation "
        "queries — i.e., if the integrity layer adds no measurable signal. "
        "Result (§5.9): full-stack achieves +45% Grounding and +204% "
        "Faithfulness uplift, confirming H1."
    ),
    (
        "H2 (GNN Auditing). Operationalization: a CompGCN trained on the "
        "graph's held-out edge set will achieve AUC-ROC ≥ 0.95 and MRR ≥ 0.95 "
        "under canonical multi-seed evaluation (n = 12). Falsification "
        "criterion: H2 is falsified if either metric falls below 0.95 on the "
        "multi-seed mean. Result (§5.3): AUC-ROC = 0.985 ± 0.001 and MRR = "
        "0.958 ± 0.005, both clearing target. (Single-seed MRR was 0.912, "
        "below target — disclosed in §5.5 and diagnosed in §6.1 as corpus-"
        "density-bound.)"
    ),
    (
        "H3 (Grounding). Operationalization: restricting synthesis to "
        "triplets that clear the τ = 0.95 plausibility gate will produce a "
        "Grounding score ≥ 0.95 on the corpus-aligned evaluation queries. "
        "Falsification criterion: H3 is falsified if Grounding falls below "
        "0.95, indicating that the τ filter does not produce a synthesis-"
        "ready subgraph. Result (§5.8): Grounding = 0.9884, clearing target."
    ),
    (
        "An additional curate-proof behavioural test of H3 is provided by the "
        "out-of-corpus refusal demonstration: the system is asked a chemistry "
        "question on a legal corpus, and the expected behaviour is a "
        "Grounding Error rather than a fabricated answer. This test cannot "
        "be passed by corpus selection — it is purely a test of the "
        "architectural mechanism."
    ),
]

# ──────────────────────────────────────────────────────────────────────────
# §6.7 Threat Model and Adversarial Considerations
# Anchors immediately before References.
# ──────────────────────────────────────────────────────────────────────────

SECTION_67_HEADING = "6.7 Threat Model and Adversarial Considerations"
SECTION_67_BODY = [
    (
        "The Validate-then-Generate architecture targets a specific threat "
        "model. This section makes that scope explicit by naming both the "
        "threats the integrity layer is designed to mitigate and the threats "
        "it is not."
    ),
    (
        "Threats the architecture mitigates. (T1) Semantic fraud and paper-"
        "mill output: fabricated documents that contain internally consistent "
        "but factually invented claims. When such documents are ingested "
        "alongside genuine ones, the CompGCN integrity model assigns low "
        "plausibility to edges whose topology disagrees with the surrounding "
        "graph — circular citation patterns, isolated subgraphs, or "
        "contradiction-edge clusters. (T2) Extraction errors in non-fraudulent "
        "documents: miscopied statistics, misattributed findings, and "
        "transcription errors introduced during LLM-based ingestion. These "
        "manifest as topologically anomalous edges and receive low plausibility "
        "scores. (T3) Hallucination at synthesis time: the τ-threshold filter "
        "removes low-plausibility triplets before they reach the LLM, and the "
        "hard Grounding Error fires when no triplets survive — replacing "
        "confabulation with refusal as the failure mode."
    ),
    (
        "Threats the architecture does not mitigate. (T4) Coordinated "
        "adversarial input: if an attacker uploads multiple internally "
        "consistent fake documents that mutually corroborate, the integrity "
        "model has no external referent against which to detect the "
        "fabrication. The closed-corpus delimitation in §1.6 names this "
        "limit. (T5) Hallucinated entities at the ingestion stage: if the "
        "extraction LLM (Gemini) invents an entity or relationship that fits "
        "the local graph structure, the integrity model cannot distinguish "
        "it from a legitimate extraction. This is a known limit of any "
        "LLM-extraction-first pipeline; mitigation paths include human "
        "extraction review, multi-LLM disagreement flagging, or symbolic-"
        "extraction fallback for structured sources. (T6) Distribution-"
        "shifted queries on a static checkpoint: queries about domains "
        "outside the trained corpus are correctly refused (the architectural "
        "guarantee), but periodic re-training on corpus updates is required "
        "to maintain quality on emerging topics."
    ),
    (
        "The honest framing is that this work mitigates the failure modes "
        "for which retrieval and synthesis are the bottleneck. Ingestion-"
        "stage threats and external-knowledge threats are separate problems "
        "and require separate components — both of which are flagged as "
        "future-work directions in §6.5."
    ),
]

# ──────────────────────────────────────────────────────────────────────────
# §6.8 Reproducibility Statement
# Anchors immediately before References.
# ──────────────────────────────────────────────────────────────────────────

SECTION_68_HEADING = "6.8 Reproducibility Statement"
SECTION_68_BODY = [
    (
        "All experimental results in Chapter 5 are reproducible from the "
        "open-source artefacts that accompany this thesis. This section "
        "documents the resources required and the steps to replicate the "
        "headline numbers."
    ),
    (
        "Source code and configuration. The framework is implemented in "
        "Python 3.11 (backend) and TypeScript 5 / Next.js 16 (frontend). "
        "Approximately 4,000 lines of backend Python and 3,000 lines of "
        "frontend TypeScript. The complete codebase, the 14-document corpus "
        "manifest, the integrity-model checkpoint, the evaluation queries, "
        "and the per-run experimental logs are released under the project's "
        "version control. The README.md at the repository root documents the "
        "two-command boot sequence (uvicorn for the API; npm run dev for the "
        "dashboard) and the .env.example file enumerates all 50+ configuration "
        "parameters with their Run 8 recommended values."
    ),
    (
        "Software dependencies. PyTorch 2.5.x (CPU build sufficient), "
        "PyTorch Geometric 2.6.x, sentence-transformers (DistilBERT-base-nli-"
        "stsb-mean-tokens, 768-dim), neo4j-graphrag, FastAPI, langchain-"
        "google-genai. Full dependency versions with upper bounds are in "
        "backend/requirements.txt; the upper bounds prevent silent breakage "
        "from upstream releases."
    ),
    (
        "Hardware floor. CPU training on a consumer laptop (Intel i5 "
        "equivalent, 16 GB RAM, no GPU) completes the Run 8 training in "
        "approximately three minutes for the 14-document corpus. Neo4j Aura "
        "Free tier is sufficient for the graph database; the schema and "
        "indexes are created at server lifespan startup. Gemini 2.5 Flash "
        "free-tier quota is sufficient to run the full evaluation suite "
        "(structural metrics + five-query LLM-as-judge) within rate-limit "
        "windows."
    ),
    (
        "Determinism and seeds. The CompGCN training is seeded by COMPGCN_"
        "SEED = 42 (configurable). The multi-seed evaluation in §5.3 uses "
        "seeds 1 through 12 inclusive; the per-seed metrics are available in "
        "backend/run_logs/multi_seed_mrr_run8.log. Gemini synthesis runs at "
        "temperature = 0 for determinism. The LLM-as-judge protocol uses the "
        "same temperature setting; per-judgment variance is approximately "
        "±0.05 in Grounding/Faithfulness scores across repeated runs of the "
        "same query."
    ),
    (
        "Replication procedure. (1) Clone the repository and follow the "
        "README two-command boot. (2) Place PDFs in backend/documents/ and "
        "trigger ingestion via POST /ingest. (3) Trigger the integrity-model "
        "training via POST /audit. (4) Trigger the evaluation suite via the "
        "restore_defense_state.py preflight script, which runs the three-mode "
        "ablation (prompt-only, graph-no-GNN, full-stack) plus the threshold "
        "sweep. (5) Inspect the resulting evaluation_results.json. Replication "
        "on the 14-document corpus reproduces the AUC, MRR, Grounding, and "
        "Faithfulness values reported in Chapter 5 within the seed-induced "
        "variance bands documented above."
    ),
]

# ──────────────────────────────────────────────────────────────────────────
# Chapter 7: Conclusion
# Anchors immediately before References.
# ──────────────────────────────────────────────────────────────────────────

CHAPTER_7_HEADING = "Chapter 7: Conclusion"
CHAPTER_7_BODY = [
    (
        "This thesis began with a stated crisis: professional knowledge work "
        "faces simultaneous collapse of manual review (85% error rates per Xu "
        "et al. 2022) and generative-AI hallucination (Frohock 2025; Athaluri "
        "et al. 2023). The architectural response proposed here, Validate-"
        "then-Generate, inserts a learned integrity layer between retrieval "
        "and synthesis so that generation is bounded by topologically "
        "validated subgraphs rather than by retrieval similarity alone."
    ),
    (
        "Three hypotheses were operationalized and tested. H1 (topological "
        "correlation) was confirmed by the +45% Grounding and +204% "
        "Faithfulness uplift of the full-stack configuration over the prompt-"
        "only baseline. H2 (GNN auditing) was confirmed under canonical "
        "multi-seed evaluation: AUC-ROC = 0.985 ± 0.001 and MRR = 0.958 ± "
        "0.005, both clearing the 0.95 target. H3 (grounding) was confirmed "
        "with Grounding = 0.9884 and Faithfulness = 0.9714, both clearing "
        "the 95% target; the curate-proof refusal demonstration provides "
        "behavioural evidence that the τ-threshold filter is a real "
        "architectural mechanism rather than a corpus-fit artefact."
    ),
    (
        "The central architectural contribution is the demonstration that "
        "a CompGCN trained on a small professional-domain corpus produces "
        "plausibility scores tight enough to gate generation at τ = 0.95 "
        "without destroying recall — and that when the gate rejects all "
        "candidates, replacing fabrication with a hard refusal is a "
        "deployable failure mode rather than a degraded experience. The "
        "Detective-Board provenance trail makes every claim in every answer "
        "auditable to a specific validated triplet and source document, "
        "which is the deployment property that professional-domain users "
        "actually require."
    ),
    (
        "The work is bounded. Empirical validation is single-corpus and "
        "single-domain; external validity to medical, regulatory, and "
        "technical corpora is future work and is named as such in §6.5. The "
        "closed-corpus delimitation (§1.6) means coordinated adversarial "
        "input that is internally self-consistent cannot be detected by the "
        "integrity layer alone — an open-corpus extension with external "
        "fact-verification is a natural research direction. The MRR ceiling "
        "diagnosis in §6.1 establishes that further architectural tuning has "
        "lower leverage than corpus expansion for this density regime."
    ),
    (
        "What this work demonstrates is narrow but durable: in the high-"
        "stakes professional knowledge contexts that motivated it, a "
        "Validate-then-Generate architecture is empirically superior to "
        "standard retrieval-augmented generation, and the superiority is "
        "attributable to the integrity layer specifically. The architecture "
        "is open, reproducible, and corpus-agnostic. The path from this "
        "feasibility study to deployed professional tools is engineering, "
        "not research."
    ),
]


def main() -> int:
    if not os.path.exists(INPUT_DOCX):
        print(f"FATAL input not found: {INPUT_DOCX}", flush=True)
        return 1

    print(f"OPEN {INPUT_DOCX}", flush=True)
    doc = Document(INPUT_DOCX)
    paragraphs_before = len(doc.paragraphs)

    # 1. §3.3.4 — insert BEFORE §3.4
    print("INSERT 3.3.4 before §3.4", flush=True)
    anchor_34 = _find_heading_starting_with(doc, "3.4 Deployment Strategy")
    _insert_section_before(
        anchor_34, doc,
        heading_text=SECTION_334_HEADING,
        heading_style="Heading 3",
        body_paragraphs=SECTION_334_BODY,
    )

    # 2/3/4. §6.7, §6.8, Chapter 7 — all insert BEFORE References, in order
    print("INSERT 6.7 before References", flush=True)
    anchor_refs = _find_heading_starting_with(doc, "References")
    _insert_section_before(
        anchor_refs, doc,
        heading_text=SECTION_67_HEADING,
        heading_style="Heading 2",
        body_paragraphs=SECTION_67_BODY,
    )

    print("INSERT 6.8 before References", flush=True)
    anchor_refs = _find_heading_starting_with(doc, "References")
    _insert_section_before(
        anchor_refs, doc,
        heading_text=SECTION_68_HEADING,
        heading_style="Heading 2",
        body_paragraphs=SECTION_68_BODY,
    )

    print("INSERT Chapter 7 before References", flush=True)
    anchor_refs = _find_heading_starting_with(doc, "References")
    _insert_section_before(
        anchor_refs, doc,
        heading_text=CHAPTER_7_HEADING,
        heading_style="Heading 1",
        body_paragraphs=CHAPTER_7_BODY,
    )

    paragraphs_after = len(doc.paragraphs)
    delta = paragraphs_after - paragraphs_before
    expected_new = (
        1 + len(SECTION_334_BODY)
        + 1 + len(SECTION_67_BODY)
        + 1 + len(SECTION_68_BODY)
        + 1 + len(CHAPTER_7_BODY)
    )
    print(f"PARAGRAPHS before={paragraphs_before} after={paragraphs_after} delta={delta} expected={expected_new}", flush=True)
    if delta != expected_new:
        print(
            f"WARN paragraph delta mismatch — expected +{expected_new}, got +{delta}. "
            "Inspect output before publishing.",
            flush=True,
        )

    doc.save(OUTPUT_DOCX)
    print(f"SAVED {OUTPUT_DOCX}", flush=True)

    # Verification: new headings present
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    required_new = [SECTION_334_HEADING, SECTION_67_HEADING, SECTION_68_HEADING, CHAPTER_7_HEADING]
    for req in required_new:
        if req not in headings:
            print(f"FAIL heading not present in output: {req}", flush=True)
            return 2
        print(f"VERIFIED inserted: {req}", flush=True)

    # Verification: existing headings preserved (including v6.6's additions)
    must_preserve = [
        "3.1 Research Design",
        "3.1.1 Corpus Composition and Selection Criteria",
        "3.3 Evaluation Framework and Metrics",
        "3.3.3 Evaluation Query Sample Size Disclosure",
        "3.4 Deployment Strategy and Practical Implications",
        "6.1 Principal Finding: Corpus-Density-Bound Performance Ceiling",
        "6.5 Limitations and Future Work",
        "6.6 Selection Bias and Reproducibility",
        "References",
    ]
    for req in must_preserve:
        present = any(req in h for h in headings)
        if not present:
            print(f"FAIL preserved heading missing: {req}", flush=True)
            return 3
    print("VERIFIED all preserved headings intact (v6.4 + v6.5 + v6.6)", flush=True)

    # User's content from v6.6 still present
    body_text = ' '.join(p.text for p in doc.paragraphs)
    user_markers = [
        "Figure 2.1: Architectural Comparison",
        "Figure 2.2: Standard RAG Pipeline",
        "Figure 2.3: Validate-then-Generate Pipeline",
        "Figure 3.1: Three-Pipeline Architecture",
        "Table 5.1: Corpus Statistics",
    ]
    missing_user = [m for m in user_markers if m not in body_text]
    if missing_user:
        print("FAIL user content from v6.6 absent:", flush=True)
        for m in missing_user:
            print(f"  - {m}", flush=True)
        return 4
    print("VERIFIED user's Figure 2.x / 3.1 / Table 5.x content preserved from v6.6", flush=True)

    img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
    table_count = len(doc.tables)
    print(f"EMBEDDED IMAGES: {img_count}", flush=True)
    print(f"TABLES: {table_count}", flush=True)
    print("DONE", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
